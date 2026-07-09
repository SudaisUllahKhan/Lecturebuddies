import streamlit as st
import requests
import PyPDF2
from docx import Document
import io
import os
import json
import base64
import time
from dotenv import load_dotenv
from document_processor import DocumentProcessor
import sounddevice as sd
import numpy as np
import soundfile as sf
import threading
import queue
import tempfile
from PIL import Image
import pytesseract
import re
import speech_recognition as sr
from faster_whisper import WhisperModel
from streamlit_mic_recorder import mic_recorder
from database import (
    create_user, 
    authenticate_user, 
    get_user_stats, 
    increment_activity,
    save_recording,
    save_quiz,
    update_user_profile,
    get_full_user_data,
    delete_user_account
)


  
   
  
    
   
   
 
    
    
  





# ==========================
import random

# ==========================
# PAGE CONFIGURATION
# ==========================
st.set_page_config(
    page_title="Lecturebuddies - Educational Platform",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

def get_random_quote():
    """Returns a random motivational quote with author"""
    quotes = [
        {"text": "The expert in anything was once a beginner.", "author": "Helen Hayes"},
        {"text": "Success is the sum of small efforts, repeated day in and day out.", "author": "Robert Collier"},
        {"text": "Live as if you were to die tomorrow. Learn as if you were to live forever.", "author": "Mahatma Gandhi"},
        {"text": "It always seems impossible until it's done.", "author": "Nelson Mandela"},
        {"text": "Don't watch the clock; do what it does. Keep going.", "author": "Sam Levenson"},
        {"text": "Education is the passport to the future, for tomorrow belongs to those who prepare for it today.", "author": "Malcolm X"},
        {"text": "The beautiful thing about learning is that no one can take it away from you.", "author": "B.B. King"},
        {"text": "Study hard, for the well is deep, and our brains are shallow.", "author": "Richard Baxter"},
        {"text": "Motivation is what gets you started. Habit is what keeps you going.", "author": "Jim Ryun"},
        {"text": "Your future is created by what you do today, not tomorrow.", "author": "Robert Kiyosaki"},
        {"text": "Believe you can and you're halfway there.", "author": "Theodore Roosevelt"},
        {"text": "Strive for progress, not perfection.", "author": "Bill Phillips"},
        {"text": "The secret of getting ahead is getting started.", "author": "Mark Twain"},
        {"text": "There are no shortcuts to any place worth going.", "author": "Beverly Sills"},
        {"text": "Focus on the goal, not the obstacle.", "author": "Anonymous"}
    ]
    return random.choice(quotes)

# ==========================
# SESSION STATE INITIALIZATION
# ==========================
def init_session_state():
    """Initialize all session state variables"""
    defaults = {
        # Login/Auth
        "authenticated": False,
        "current_user": None,
        "user_id": None, 
        "user_stats": {}, 
        "current_page": "dashboard",

        # Chatbot & Summarization
        "messages": [],
        "uploaded_files": [],
        "document_contents": {},
        "chat_model": "llama-3.1-8b-instant",
        "chat_temperature": 0.7,

        # Quiz Generator
        "quiz_output": None,
        "num_questions": 5,
        "difficulty": "Medium",
        "quiz_model": "llama-3.1-8b-instant",
        "quiz_temperature": 0.7,

        # Live Lecture Recording
        "rec_thread": None,
        "audio_queue": None,
        "recording": False,
        "transcript": "",
        "partial_transcript": "",
        "chunks_saved": [],

        # Dashboard
        "selected_feature": None
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_session_state()

# ==========================
# LOAD API KEY
# ==========================
load_dotenv()
api_key = os.getenv("GROQ_API_KEY")

# ==========================
# LOGO HELPER
# ==========================
def get_logo_svg(size_px=80, font_size_px=42):
    """Returns the SVG logo string with adjustable size wrapped in a link"""
    svg_icon = f'<svg class="lb-logo-svg" width="{size_px}" height="{size_px}" viewBox="0 0 100 100" fill="none" xmlns="http://www.w3.org/2000/svg" style="position: relative; top: 10px;"><path d="M50 15L20 28L50 41L80 28L50 15Z" fill="#2D005E"/><path d="M70 32V45" stroke="#2D005E" stroke-width="2"/><circle cx="50" cy="48" r="15" fill="#2D005E"/></svg>'
    font_style = f'font-family: \'Georgia\', serif; font-size: {font_size_px}px; font-weight: 900; color: #2D005E; letter-spacing: -1px; margin-left: -10px;'
    
    return f'<a href="https://lecturebuddies.streamlit.app" target="_blank" style="text-decoration: none; display: flex; align-items: center; gap: 0px; cursor: pointer;">{svg_icon}<span class="lb-logo-text" style="{font_style}">Lecturebuddies</span></a>'

# ==========================
# GLOBAL STYLING - LECTUREBUDDIES THEME
# ==========================
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

    /* ==============================================
       LECTUREBUDDIES — CLEAN SAAS DESIGN SYSTEM
       White surfaces, thin borders, crisp sans-serif,
       purple (#4e54c8) used only as an accent.
       ============================================== */
    :root {
        --ink: #16182d;
        --text: #3f4257;
        --muted: #6f7287;
        --faint: #9becac;
        --border: #e7e8f0;
        --border-hover: #d5d7e6;
        --bg: #f7f8fa;
        --surface: #ffffff;
        --primary: #4e54c8;
        --primary-hover: #4046b8;
        --primary-soft: #eef0fd;
        --primary-border: #d9dcf7;
        --radius: 10px;
        --radius-lg: 12px;
        --shadow-xs: 0 1px 2px rgba(22, 24, 45, 0.05);
        --shadow-sm: 0 1px 3px rgba(22, 24, 45, 0.06), 0 4px 12px rgba(22, 24, 45, 0.04);
        --shadow-md: 0 4px 16px rgba(22, 24, 45, 0.08);
        --font: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    }

    /* Base canvas — flat, quiet */
    .stApp {
        background: var(--bg);
        color: var(--text);
    }

    /* Professional sans-serif everywhere (overrides legacy inline Georgia)… */
    .stApp * { font-family: var(--font) !important; }
    /* …except the wordmark, which stays serif for brand */
    .lb-logo-text { font-family: Georgia, 'Times New Roman', serif !important; }
    /* …and Streamlit's Material icon glyphs, which are font ligatures —
       forcing Inter on them shows raw names like "expand_more" */
    .stApp [data-testid="stIconMaterial"],
    .stApp [data-testid="stExpanderToggleIcon"],
    .stApp span[translate="no"],
    .stApp [class*="material-symbols"] {
        font-family: 'Material Symbols Rounded' !important;
    }

    /* Hide Streamlit chrome, keep sidebar toggle on mobile */
    #MainMenu, footer, [data-testid="stAppToolbar"], .stDeployButton {
        visibility: hidden !important;
        display: none !important;
    }
    [data-testid="stHeader"] { background-color: transparent !important; }
    @media (max-width: 768px) {
        [data-testid="collapsedControl"] {
            visibility: visible !important;
            display: flex !important;
            background: #ffffff !important;
            border: 1px solid var(--border) !important;
            border-radius: 8px !important;
            margin: 10px !important;
            box-shadow: var(--shadow-xs) !important;
        }
    }
    @media (min-width: 769px) {
        [data-testid="stHeader"] { display: none !important; }
        [data-testid="collapsedControl"] { display: none !important; }
        [data-testid="stSidebar"][aria-expanded="true"],
        [data-testid="stSidebar"][aria-expanded="false"] {
            transform: none !important;
            margin-left: 0 !important;
        }
        section[data-testid="stSidebar"] {
            position: relative !important;
            min-width: 17.5rem !important;
            max-width: 17.5rem !important;
        }
    }

    section[data-testid="stMain"] .block-container {
        padding: 0.9rem clamp(1rem, 3vw, 2.4rem) 2.4rem clamp(1rem, 3vw, 2.4rem) !important;
        max-width: 1200px;
        margin: 0 auto;
    }
    /* Tighter vertical rhythm between elements */
    section[data-testid="stMain"] div[data-testid="stVerticalBlock"] { gap: 0.75rem !important; }

    /* ==============================================
       TYPOGRAPHY
       ============================================== */
    section[data-testid="stMain"] h1, section[data-testid="stMain"] h2,
    section[data-testid="stMain"] h3, section[data-testid="stMain"] h4 {
        color: var(--ink);
        letter-spacing: -0.02em;
        font-weight: 700;
    }
    section[data-testid="stMain"] p, section[data-testid="stMain"] li { color: var(--text); }

    .main-title {
        text-align: center;
        font-size: clamp(26px, 4.5vw, 36px);
        font-weight: 800;
        color: var(--ink);
        letter-spacing: -0.03em;
        margin-bottom: 6px;
    }
    .tagline {
        text-align: center;
        font-size: clamp(13px, 2vw, 15px);
        color: var(--muted);
        margin-bottom: 15px;
    }
    .app-brand-title {
        font-size: clamp(26px, 4.5vw, 36px);
        font-weight: 800;
        color: var(--ink);
        letter-spacing: -0.03em;
        margin-bottom: 8px;
        text-align: center;
    }
    hr.gradient {
        border: none;
        height: 2px;
        background: linear-gradient(90deg, #4e54c8, #8f94fb, #764ba2);
        margin: 2px auto 16px auto;
        max-width: 100%;
        border-radius: 2px;
        opacity: 0.9;
    }

    /* ==============================================
       PAGE HEADER (welcome-container, reused everywhere)
       — clean left-aligned header card, no gradients
       ============================================== */
    .welcome-container {
        text-align: left;
        padding: clamp(22px, 3.2vw, 32px) clamp(20px, 3vw, 34px);
        background: linear-gradient(120deg, #4e54c8 0%, #667eea 55%, #764ba2 100%);
        border: none;
        border-radius: var(--radius-lg);
        margin: 4px 0 20px 0;
        box-shadow: 0 8px 24px rgba(78, 84, 200, 0.25);
    }
    .welcome-eyebrow {
        display: inline-block;
        color: #ffffff !important;
        background: rgba(255, 255, 255, 0.16);
        border: 1px solid rgba(255, 255, 255, 0.28);
        border-radius: 6px;
        padding: 3px 10px;
        font-size: 11px;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.8px;
        margin-bottom: 12px;
    }
    .welcome-title {
        color: #ffffff !important;
        font-size: clamp(20px, 3vw, 27px);
        font-weight: 700;
        letter-spacing: -0.02em;
        margin: 0 0 6px 0;
    }
    .welcome-text {
        color: rgba(255, 255, 255, 0.90) !important;
        font-size: clamp(13.5px, 1.8vw, 15px);
        margin: 0;
        line-height: 1.6;
        max-width: 760px;
    }

    /* ==============================================
       SIDEBAR — light, quiet, professional
       ============================================== */
    [data-testid="stSidebar"] {
        background: #ffffff;
        border-right: 1px solid var(--border);
    }
    [data-testid="stSidebar"] > div:first-child {
        padding-top: 14px !important;
        padding-bottom: 14px !important;
    }
    [data-testid="stSidebar"] p, [data-testid="stSidebar"] li,
    [data-testid="stSidebar"] label, [data-testid="stSidebar"] label p {
        color: var(--text) !important;
    }
    [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3, [data-testid="stSidebar"] h4 {
        color: var(--ink) !important;
        font-size: 14px !important;
        font-weight: 600 !important;
        letter-spacing: -0.01em;
    }
    [data-testid="stSidebar"] hr {
        border: none !important;
        height: 1px !important;
        background: var(--border) !important;
        margin: 12px 0 !important;
    }
    .sidebar-title {
        font-size: 11px;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 1px;
        color: var(--muted);
        padding: 0 12px;
        margin-bottom: 8px;
    }

    /* Sidebar nav buttons — quiet rows, subtle hover */
    [data-testid="stSidebar"] button {
        width: 100% !important;
        text-align: left !important;
        justify-content: flex-start !important;
        padding: 9px 12px !important;
        margin: 1px 0 !important;
        border-radius: 8px !important;
        background: transparent !important;
        border: none !important;
        color: #4b5563 !important;
        font-weight: 500 !important;
        font-size: 13.5px !important;
        transition: background 0.15s ease, color 0.15s ease !important;
        display: flex !important;
        align-items: center !important;
        gap: 8px !important;
        white-space: nowrap !important;
        min-height: 36px !important;
        box-shadow: none !important;
    }
    [data-testid="stSidebar"] button p { color: inherit !important; font-size: 13.5px !important; }
    [data-testid="stSidebar"] button:hover {
        background: var(--primary-soft) !important;
        color: var(--primary) !important;
        transform: none;
        box-shadow: none !important;
    }
    [data-testid="stSidebar"] button:active { background: #e3e6fb !important; }

    /* Sidebar expander — invisible frame, reads as a plain menu */
    [data-testid="stSidebar"] [data-testid="stExpander"] {
        border: none !important;
        border-radius: 0 !important;
        background: transparent !important;
        box-shadow: none !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] details {
        border: none !important;
        background: transparent !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] summary {
        font-weight: 600 !important;
        font-size: 12px !important;
        text-transform: uppercase;
        letter-spacing: 0.8px;
        color: var(--muted) !important;
        padding: 2px 8px !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] summary:hover { color: var(--ink) !important; }

    /* ==============================================
       BUTTONS — flat, solid, professional
       ============================================== */
    .stButton button {
        width: 100% !important;
        background: linear-gradient(135deg, #4e54c8 0%, #667eea 55%, #764ba2 100%) !important;
        color: #ffffff !important;
        font-weight: 600 !important;
        font-size: 14px !important;
        padding: 9px 16px !important;
        border: none !important;
        border-radius: 8px !important;
        cursor: pointer !important;
        transition: filter 0.15s ease, box-shadow 0.15s ease !important;
        box-shadow: 0 2px 8px rgba(78, 84, 200, 0.28) !important;
        margin-top: 6px !important;
    }
    .stButton button:hover {
        filter: brightness(1.08);
        box-shadow: 0 4px 14px rgba(78, 84, 200, 0.38) !important;
        transform: none !important;
    }
    .stButton button:focus-visible {
        outline: 2px solid var(--primary) !important;
        outline-offset: 2px !important;
    }
    .stButton button[kind="primary"] {
        background: linear-gradient(135deg, #4e54c8 0%, #667eea 55%, #764ba2 100%) !important;
        border: none !important;
        color: #ffffff !important;
    }
    .stButton button[kind="secondary"] {
        background: #ffffff !important;
        color: var(--text) !important;
        border: 1px solid var(--border) !important;
        box-shadow: var(--shadow-xs) !important;
    }
    .stButton button[kind="secondary"]:hover {
        border-color: var(--border-hover) !important;
        background: #fafafc !important;
        color: var(--ink) !important;
    }

    /* Download buttons — outlined */
    .stDownloadButton button {
        width: 100% !important;
        background: #ffffff !important;
        color: var(--primary) !important;
        border: 1px solid var(--primary-border) !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        font-size: 14px !important;
        padding: 9px 16px !important;
        box-shadow: var(--shadow-xs) !important;
        transition: all 0.15s ease !important;
    }
    .stDownloadButton button:hover {
        background: var(--primary-soft) !important;
        border-color: var(--primary) !important;
        transform: none !important;
    }

    /* ==============================================
       INPUTS / TEXTAREA / SELECT
       ============================================== */
    .stTextInput input { color: var(--ink) !important; }
    ::placeholder { color: #a3a6bd !important; opacity: 1 !important; }
    .stTextInput input::placeholder,
    .stTextArea textarea::placeholder { color: #a3a6bd !important; font-weight: 400 !important; }

    .stTextInput > div > div > input,
    .stTextInput > div > input {
        border: 1px solid var(--border) !important;
        border-radius: 8px !important;
        padding: 10px 12px !important;
        font-size: 14px !important;
        transition: border-color 0.15s ease, box-shadow 0.15s ease !important;
        height: auto !important;
        width: 100% !important;
        background: #ffffff !important;
        box-shadow: var(--shadow-xs) !important;
    }
    .stTextInput > div > div > input:focus,
    .stTextInput > div > input:focus {
        border-color: var(--primary) !important;
        outline: none !important;
        box-shadow: 0 0 0 3px rgba(78, 84, 200, 0.12) !important;
    }
    .stTextInput div[data-baseweb="input"] {
        border-radius: 8px !important;
        background: transparent !important;
        border: none !important;
    }

    .stTextArea textarea {
        border: 1px solid var(--border) !important;
        border-radius: 8px !important;
        padding: 12px !important;
        font-size: 14px !important;
        background: #ffffff !important;
        color: var(--ink) !important;
        transition: border-color 0.15s ease, box-shadow 0.15s ease !important;
        box-shadow: var(--shadow-xs) !important;
    }
    .stTextArea textarea:focus {
        border-color: var(--primary) !important;
        box-shadow: 0 0 0 3px rgba(78, 84, 200, 0.12) !important;
        outline: none !important;
    }
    .stTextArea div[data-baseweb="textarea"] {
        border-radius: 8px !important;
        border: none !important;
        background: transparent !important;
    }

    /* Labels */
    .stTextInput label, .stTextInput label p,
    .stTextArea label, .stTextArea label p,
    .stSelectbox label, .stSelectbox label p,
    .stMultiSelect label, .stMultiSelect label p,
    .stDateInput label, .stDateInput label p,
    .stFileUploader label, .stFileUploader label p,
    .stSlider label, .stCheckbox label p {
        color: var(--ink) !important;
        font-weight: 500 !important;
        font-size: 13.5px !important;
    }

    /* Selectbox / Multiselect */
    .stSelectbox div[data-baseweb="select"] > div,
    .stMultiSelect div[data-baseweb="select"] > div {
        border: 1px solid var(--border) !important;
        border-radius: 8px !important;
        background: #ffffff !important;
        transition: border-color 0.15s ease !important;
        min-height: 40px;
        box-shadow: var(--shadow-xs);
    }
    .stSelectbox div[data-baseweb="select"] > div:hover,
    .stMultiSelect div[data-baseweb="select"] > div:hover {
        border-color: var(--border-hover) !important;
    }
    div[data-baseweb="popover"] ul[role="listbox"] {
        border-radius: 8px !important;
        border: 1px solid var(--border) !important;
        box-shadow: var(--shadow-md) !important;
    }
    .stMultiSelect span[data-baseweb="tag"] {
        background: var(--primary-soft) !important;
        border: 1px solid var(--primary-border) !important;
        border-radius: 6px !important;
        color: var(--primary) !important;
    }
    .stMultiSelect span[data-baseweb="tag"] span { color: var(--primary) !important; }

    .stDateInput div[data-baseweb="input"] {
        border: 1px solid var(--border) !important;
        border-radius: 8px !important;
        background: #ffffff !important;
    }

    /* ==============================================
       TABS — underline style (Linear/Stripe)
       ============================================== */
    .stTabs [data-baseweb="tab-list"] {
        gap: 2px;
        background: transparent;
        border-radius: 0;
        padding: 0;
        margin-bottom: 22px;
        border-bottom: 1px solid var(--border);
        flex-wrap: wrap;
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: 6px 6px 0 0;
        font-weight: 500;
        font-size: 14px;
        padding: 9px 14px;
        transition: color 0.15s ease, background 0.15s ease;
        border: none;
        cursor: pointer;
        color: var(--muted);
        background: transparent;
        margin-bottom: 0;
    }
    .stTabs [data-baseweb="tab"] p { font-size: 14px !important; color: inherit !important; }
    .stTabs [aria-selected="true"] {
        background: transparent !important;
        color: var(--primary) !important;
        box-shadow: none !important;
        font-weight: 600;
    }
    .stTabs [aria-selected="true"] p { color: var(--primary) !important; }
    .stTabs [data-baseweb="tab"]:hover:not([aria-selected="true"]) {
        background: #f2f3f9;
        color: var(--ink);
    }
    .stTabs [data-baseweb="tab-highlight"] {
        background-color: var(--primary) !important;
        height: 2px !important;
    }
    .stTabs [data-baseweb="tab-border"] { display: none !important; }

    /* ==============================================
       FILE UPLOADER
       ============================================== */
    [data-testid="stFileUploaderDropzone"] {
        background: #fafafc !important;
        border: 1px dashed var(--border-hover) !important;
        border-radius: var(--radius) !important;
        transition: border-color 0.15s ease, background 0.15s ease !important;
    }
    [data-testid="stFileUploaderDropzone"]:hover {
        border-color: var(--primary) !important;
        background: var(--primary-soft) !important;
    }
    [data-testid="stFileUploaderDropzone"] button {
        background: #ffffff !important;
        color: var(--text) !important;
        border: 1px solid var(--border) !important;
        border-radius: 8px !important;
        font-weight: 500 !important;
        transition: border-color 0.15s ease !important;
    }
    [data-testid="stFileUploaderDropzone"] button:hover {
        border-color: var(--border-hover) !important;
    }
    [data-testid="stFileUploaderDropzone"] span,
    [data-testid="stFileUploaderDropzone"] small { color: var(--muted) !important; }

    /* ==============================================
       EXPANDERS
       ============================================== */
    section[data-testid="stMain"] [data-testid="stExpander"] {
        border: 1px solid var(--border) !important;
        border-radius: var(--radius) !important;
        background: #ffffff !important;
        box-shadow: var(--shadow-xs) !important;
        overflow: hidden;
        margin-bottom: 8px;
    }
    section[data-testid="stMain"] [data-testid="stExpander"] summary {
        font-weight: 500 !important;
        color: var(--text) !important;
        padding: 4px 6px;
        font-size: 14px !important;
    }
    section[data-testid="stMain"] [data-testid="stExpander"] summary:hover { color: var(--primary) !important; }

    /* ==============================================
       METRICS
       ============================================== */
    [data-testid="stMetric"] {
        background: #ffffff;
        border: 1px solid var(--border);
        border-radius: var(--radius);
        padding: 14px 16px;
        box-shadow: var(--shadow-xs);
    }
    [data-testid="stMetricValue"] {
        color: var(--ink) !important;
        font-weight: 700 !important;
    }
    [data-testid="stMetricLabel"] { color: var(--muted) !important; }

    /* ==============================================
       ALERTS
       ============================================== */
    [data-testid="stAlert"] {
        border-radius: 8px !important;
        border: 1px solid var(--border) !important;
        box-shadow: none !important;
        font-size: 14px !important;
    }

    /* ==============================================
       SLIDERS
       ============================================== */
    .stSlider [data-baseweb="slider"] div[role="slider"] {
        background-color: #ffffff !important;
        border: 2px solid var(--primary) !important;
        box-shadow: var(--shadow-xs) !important;
    }
    .stSlider > div > div > div > div { background-color: var(--primary) !important; }
    .stSlider > div > div > div { background-color: #e3e4f0 !important; }
    .stSlider [data-testid="stTickBarMin"],
    .stSlider [data-testid="stTickBarMax"] { color: var(--muted) !important; }
    .stSlider div[data-testid="stSliderThumbValue"] {
        color: var(--primary) !important;
        font-weight: 600 !important;
    }

    /* ==============================================
       CHAT
       ============================================== */
    .user-message { text-align: right; margin: 8px 0; }
    .assistant-message { margin: 8px 0; }
    .user-bubble {
        display: inline-block;
        background: linear-gradient(135deg, #4e54c8, #667eea);
        color: #ffffff;
        padding: 10px 14px;
        border-radius: 14px 14px 4px 14px;
        max-width: 72%;
        font-size: 14px;
        line-height: 1.55;
        text-align: left;
        word-wrap: break-word;
    }
    .assistant-bubble {
        display: inline-block;
        background: #ffffff;
        padding: 10px 14px;
        border-radius: 14px 14px 14px 4px;
        max-width: 78%;
        border: 1px solid var(--border);
        font-size: 14px;
        line-height: 1.6;
        word-wrap: break-word;
        color: var(--text);
        box-shadow: var(--shadow-xs);
    }

    [data-testid="stChatInput"] {
        border: 1px solid var(--border) !important;
        border-radius: 10px !important;
        background: #ffffff !important;
        box-shadow: var(--shadow-xs) !important;
        transition: border-color 0.15s ease, box-shadow 0.15s ease !important;
    }
    [data-testid="stChatInput"]:focus-within {
        border-color: var(--primary) !important;
        box-shadow: 0 0 0 3px rgba(78, 84, 200, 0.12) !important;
    }
    [data-testid="stChatInput"] textarea { color: var(--ink) !important; }
    [data-testid="stChatInput"] button {
        background: var(--primary) !important;
        border-radius: 8px !important;
        color: white !important;
    }

    /* ==============================================
       POPOVER (profile menu)
       ============================================== */
    [data-testid="stPopoverBody"] {
        border-radius: var(--radius) !important;
        border: 1px solid var(--border) !important;
        box-shadow: var(--shadow-md) !important;
    }

    /* ==============================================
       STAT CARDS
       ============================================== */
    .stats-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
        gap: 14px;
        margin-bottom: 24px;
    }
    .stat-card {
        background: #ffffff;
        padding: 18px 20px;
        border-radius: var(--radius);
        border: 1px solid var(--border);
        box-shadow: var(--shadow-xs);
        transition: border-color 0.15s ease, box-shadow 0.15s ease;
    }
    .stat-card:hover {
        border-color: var(--border-hover);
        box-shadow: var(--shadow-sm);
    }
    .stat-label {
        color: var(--muted);
        font-size: 12px;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.6px;
        margin: 0 0 8px 0;
    }
    .stat-value {
        background: linear-gradient(120deg, #4e54c8, #764ba2);
        -webkit-background-clip: text;
        background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 28px;
        font-weight: 700;
        letter-spacing: -0.02em;
        margin: 0 0 4px 0;
        line-height: 1.1;
    }
    .stat-card { border-top: 3px solid #667eea; }
    .stat-sub {
        color: var(--muted);
        font-size: 12.5px;
        margin: 0;
    }
    /* Legacy classes kept for safety */
    .feature-title { color: var(--ink) !important; font-size: 15px !important; font-weight: 600 !important; margin: 0 0 4px 0 !important; }
    .feature-description { color: var(--muted) !important; font-size: 13px !important; margin: 0 !important; }

    /* Quote card */
    .quote-card {
        text-align: center;
        margin-top: 20px;
        background: #ffffff;
        padding: clamp(22px, 3vw, 32px);
        border: 1px solid var(--border);
        border-radius: var(--radius-lg);
        box-shadow: var(--shadow-xs);
    }
    .quote-card h3 {
        color: var(--muted);
        font-size: 12px;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.8px;
        margin: 0 0 12px 0;
    }
    .quote-card .quote-text {
        color: var(--ink);
        font-size: clamp(15px, 2vw, 17px);
        font-style: italic;
        margin: 0 0 8px 0;
        line-height: 1.6;
    }
    .quote-card .quote-author {
        color: var(--muted);
        font-size: 13px;
        margin: 6px 0 0 0;
    }

    /* ==============================================
       QUIZ RESULTS
       ============================================== */
    .quiz-container {
        background: #ffffff;
        border: 1px solid var(--border);
        border-radius: var(--radius-lg);
        padding: 20px 22px 12px 22px;
        box-shadow: var(--shadow-xs);
        margin: 16px 0 12px 0;
    }

    /* ==============================================
       FLASH CARDS
       ============================================== */
    .flash-card {
        background: #ffffff;
        padding: clamp(20px, 3vw, 28px);
        border-radius: var(--radius-lg);
        border: 1px solid var(--border);
        box-shadow: var(--shadow-sm);
        text-align: center;
        margin: 18px 0;
    }
    .flash-card-counter {
        display: inline-block;
        background: var(--primary-soft);
        color: var(--primary);
        border: 1px solid var(--primary-border);
        font-size: 12px;
        font-weight: 600;
        padding: 4px 12px;
        border-radius: 6px;
        margin-bottom: 14px;
    }
    .flash-card-face {
        background: #fafafc;
        border: 1px solid var(--border);
        padding: clamp(16px, 2.5vw, 24px);
        border-radius: var(--radius);
        margin: 6px 0;
    }
    .flash-card-face h4 {
        color: var(--muted);
        margin: 0 0 8px 0;
        font-size: 11px;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    .flash-card-face p {
        font-size: clamp(15px, 2.2vw, 17px);
        color: var(--ink);
        margin: 0;
        line-height: 1.55;
    }
    .flash-card-face.answer {
        background: #f4faf6;
        border-color: #d3ecd9;
    }
    .flash-card-face.answer h4 { color: #1f9d55; }

    /* ==============================================
       COMING SOON
       ============================================== */
    .coming-soon {
        text-align: center;
        background: #ffffff;
        border: 1px solid var(--border);
        border-radius: var(--radius-lg);
        padding: clamp(40px, 6vw, 64px) 24px;
        box-shadow: var(--shadow-xs);
        margin-top: 16px;
    }
    .coming-soon-icon { font-size: 44px; margin-bottom: 14px; }
    .coming-soon-title { color: var(--ink); font-weight: 700; margin-bottom: 10px; }
    .coming-soon-text {
        color: var(--muted);
        max-width: 480px;
        margin: 0 auto;
        line-height: 1.6;
        font-size: 14px;
    }

    /* ==============================================
       FOOTER
       ============================================== */
    .lb-footer {
        text-align: center;
        color: #a3a6bd;
        font-size: 12px;
        margin: 8px 0 0 0;
    }
    .app-footer {
        display: flex;
        justify-content: space-between;
        align-items: center;
        flex-wrap: wrap;
        gap: 6px 16px;
        margin-top: 44px;
        padding: 16px 4px 4px 4px;
        border-top: 1px solid var(--border);
        color: #8b8ea6;
        font-size: 12.5px;
    }
    .app-footer .footer-brand { font-weight: 600; color: var(--text); }
    .app-footer .footer-credit b {
        font-weight: 600;
        background: linear-gradient(120deg, #4e54c8, #764ba2);
        -webkit-background-clip: text;
        background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    @media (max-width: 600px) {
        .app-footer { justify-content: center; text-align: center; }
    }

    .dev-access-container {
        margin-top: 30px;
        padding: 20px;
        background: #fff7f7;
        border: 1px dashed #f3c1c1;
        border-radius: var(--radius);
        text-align: center;
    }

    /* Login extras */
    .login-links {
        text-align: center;
        margin-top: 22px;
        padding-top: 22px;
        border-top: 1px solid var(--border);
    }
    .login-link {
        color: var(--primary);
        text-decoration: none;
        font-size: 13.5px;
        font-weight: 500;
    }
    .login-link:hover { text-decoration: underline; }

    /* ==============================================
       SCROLLBARS
       ============================================== */
    ::-webkit-scrollbar { width: 10px; height: 10px; }
    ::-webkit-scrollbar-track { background: transparent; }
    ::-webkit-scrollbar-thumb {
        background: #d5d7e6;
        border-radius: 999px;
        border: 2px solid transparent;
        background-clip: content-box;
    }
    ::-webkit-scrollbar-thumb:hover { background-color: #b8bbd4; }

    /* ==============================================
       RESPONSIVE
       ============================================== */
    @media only screen and (max-width: 768px) {
        section[data-testid="stMain"] .block-container { padding: 1rem 0.9rem 2rem 0.9rem !important; }
        .user-bubble, .assistant-bubble {
            max-width: 92% !important;
            font-size: 13.5px !important;
        }
        [data-testid="stSidebar"] { min-width: 260px !important; }
        .stTabs [data-baseweb="tab"] { padding: 8px 10px; font-size: 13.5px; }
    }
    @media only screen and (max-width: 480px) {
        section[data-testid="stMain"] .block-container { padding: 0.75rem 0.65rem 2rem 0.65rem !important; }
        .stats-grid { grid-template-columns: 1fr 1fr; gap: 10px; }
        .stat-value { font-size: 22px; }
        .login-logo-wrapper .lb-logo-svg {
            width: 56px !important;
            height: 56px !important;
            min-width: 56px !important;
        }
        .login-logo-wrapper .lb-logo-text {
            font-size: 26px !important;
            margin-left: -5px !important;
        }
    }
    @media only screen and (min-width: 1600px) {
        section[data-testid="stMain"] .block-container { max-width: 1320px; }
        .user-bubble, .assistant-bubble { max-width: 70% !important; }
    }

    @media (prefers-reduced-motion: reduce) {
        * { transition: none !important; animation: none !important; }
    }
    </style>
    """,
    unsafe_allow_html=True
)

# ==========================
# SIMPLE PASSWORD RESET
# ==========================
def show_simple_password_reset():
    """Simple password reset using only username"""
    st.markdown("""
        <div class="welcome-container">
            <h1 class="welcome-title">🔐 Reset Password</h1>
            <p class="welcome-text">Enter your username to set a new password</p>
        </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 1.5, 1])
    with col2:
        reset_username = st.text_input("Username", placeholder="Enter your username", key="reset_user_input")
        
        if reset_username:
            from database import get_user_id_by_username, update_user_profile
            user_id = get_user_id_by_username(reset_username)
            
            if user_id:
                st.success(f"User '{reset_username}' found. Enter your new password below.")
                new_pass = st.text_input("New Password", type="password", key="new_pass_simple")
                confirm_pass = st.text_input("Confirm New Password", type="password", key="confirm_pass_simple")
                
                if st.button("Update Password", use_container_width=True):
                    if new_pass == confirm_pass:
                        if len(new_pass) < 6:
                            st.error("Password must be at least 6 characters long")
                        else:
                            success, message = update_user_profile(user_id, new_password=new_pass)
                            if success:
                                st.success("Password updated successfully!")
                                time.sleep(1.5)
                                st.session_state.show_simple_reset = False
                                st.rerun()
                            else:
                                st.error(message)
                    else:
                        st.error("Passwords do not match")
            else:
                st.error("User not found")
        
        if st.button("← Back to Login", use_container_width=True):
            st.session_state.show_simple_reset = False
            st.rerun()

# ==========================
# APP FOOTER (display only)
# ==========================
def render_app_footer():
    """Professional footer shown at the bottom of every page"""
    st.markdown(
        """
        <div class="app-footer">
            <span><span class="footer-brand">© 2026 LectureBuddies.</span> All rights reserved.</span>
            <span class="footer-credit">Designed &amp; developed by <b>Autovex Solutions</b></span>
        </div>
        """,
        unsafe_allow_html=True
    )

# ==========================
# LOGIN AND SIGNUP SECTION
# ==========================
def show_login_page():
    """Display login/signup page with WordPress-style centered form"""
    
    # --- CSS INJECTION FOR BOLD FONTS ---
    st.markdown(
        """
        <style>
        /* ===== AUTH PAGE ===== */
        section[data-testid="stMain"] .block-container {
            padding-top: clamp(1.5rem, 6vh, 3.5rem) !important;
            max-width: 1080px;
        }

        /* Login / Sign Up switcher spans the column */
        .stTabs [data-baseweb="tab-list"] { justify-content: center; }
        .stTabs [data-baseweb="tab"] {
            flex: 1;
            justify-content: center;
            text-align: center;
        }

        /* Primary auth buttons slightly larger */
        .st-key-login_btn button, .st-key-signup_btn button {
            padding: 11px 16px !important;
            font-size: 14.5px !important;
        }

        /* "Forgot Password?" as a quiet text link */
        .st-key-forgot_pass_simple_btn button {
            background: transparent !important;
            color: #4e54c8 !important;
            box-shadow: none !important;
            border: none !important;
            font-weight: 500 !important;
            font-size: 13px !important;
            padding: 4px 8px !important;
            margin-top: 0 !important;
            width: auto !important;
        }
        .st-key-forgot_pass_simple_btn button:hover {
            text-decoration: underline;
            background: transparent !important;
        }
        .st-key-forgot_pass_simple_btn { display: flex; justify-content: center; }
        </style>
        """,
        unsafe_allow_html=True
    )
    # ------------------------------------

    # Center column approach for login form
    col1, col2, col3 = st.columns([1, 1.5, 1])
    
    with col2:
        # Display SVG logo and tagline
        logo_html = get_logo_svg(size_px=100, font_size_px=48)
        tagline_html = '<div style="text-align: center; margin-top: 5px; margin-bottom: 30px;"><p style="color: #666; font-size: 16px; font-family: \'Georgia\', serif; margin-top: 5px; line-height: 1.6;">Your intelligent study companion—learn, create, and excel with AI</p></div>'
        st.markdown(f'<div class="login-logo-wrapper" style="display: flex; flex-direction: column; align-items: center; margin-bottom: 20px;">{logo_html}{tagline_html}</div>', unsafe_allow_html=True)
        
       
        
        # Login/Signup Tabs
        tab1, tab2 = st.tabs(["Login", "Sign Up"])

        with tab1:
            username = st.text_input("Username", placeholder="Enter your username", key="login_username")
            password = st.text_input("Password", type="password", placeholder="Enter your password", key="login_password")

            if st.button("Login", key="login_btn", help="Login to your account"):
                if username and password:
                    # Authenticate using database
                    success, message, user_id = authenticate_user(username, password)
                    if success:
                        st.session_state.authenticated = True
                        st.session_state.current_user = username
                        st.session_state.user_id = user_id
                        # Load user stats
                        st.session_state.user_stats = get_user_stats(user_id)
                        st.success(message)
                        st.rerun()
                    else:
                        st.error(message)
                else:
                    st.warning("Please enter both username and password")
            
            # Simple Forgot Password Link
            st.markdown('<div style="text-align: center; margin-top: 10px;">', unsafe_allow_html=True)
            if st.button("Forgot Password?", key="forgot_pass_simple_btn"):
                st.session_state.show_simple_reset = True
                st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
            

        with tab2:
            new_username = st.text_input("Choose Username", placeholder="Enter your username", key="signup_username")
            new_email = st.text_input("Email (optional)", placeholder="Enter your email", key="signup_email")
            new_password = st.text_input("Choose Password", type="password", placeholder="Enter your password", key="signup_password")
            confirm_password = st.text_input("Confirm Password", type="password", placeholder="Confirm your password", key="confirm_password")

            if st.button("Create Account", key="signup_btn", help="Create new account"):
                if new_username and new_password and confirm_password:
                    if new_password == confirm_password:
                        if len(new_password) < 6:
                            st.error("Password must be at least 6 characters long")
                        else:
                            # Create user in database
                            success, message, user_id = create_user(new_username, new_password, new_email if new_email else None)
                            if success:
                                st.session_state.authenticated = True
                                st.session_state.current_user = new_username
                                st.session_state.user_id = user_id
                                # Initialize user stats (will be all zeros for new user)
                                st.session_state.user_stats = get_user_stats(user_id)
                                st.success(message)
                                st.rerun()
                            else:
                                st.error(message)
                    else:
                        st.error("Passwords don't match")
                else:
                    st.warning("Please fill in username and password fields")

        # Footer under the auth form
        render_app_footer()



# ==========================
# DASHBOARD LAYOUT SECTION
# ==========================
def show_dashboard():
    """Display main dashboard with sidebar and content area"""
    
    # Initialize profile dropdown state
    if "show_profile_dropdown" not in st.session_state:
        st.session_state.show_profile_dropdown = False

    # Add JavaScript to detect sidebar state and show/hide floating toggle button


    # --- RESPONSIVE TOP NAVIGATION BAR ---
    with st.container():
        # CSS for horizontal alignment and styling
        st.markdown("""
            <style>
                .nav-container {
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                    padding: 0.5rem 1rem;
                    background-color: transparent;
                    margin-bottom: 1rem;
                }
                /* User profile button — quiet outlined control */
                [data-testid="stPopover"] > button {
                    background: #ffffff !important;
                    border: 1px solid #e7e8f0 !important;
                    color: #3f4257 !important;
                    border-radius: 8px !important;
                    padding: 0.45rem 0.9rem !important;
                    box-shadow: 0 1px 2px rgba(22, 24, 45, 0.05) !important;
                    font-weight: 500 !important;
                    height: auto !important;
                    width: auto !important;
                    transition: border-color 0.15s ease, background 0.15s ease !important;
                }
                [data-testid="stPopover"] > button:hover {
                    border-color: #d5d7e6 !important;
                    background: #fafafc !important;
                    transform: none !important;
                }
                [data-testid="stPopover"] > button p {
                    font-size: 14px !important;
                    font-weight: 500 !important;
                    color: #3f4257 !important;
                }
                /* Keep navbar row vertically aligned */
                div[data-testid="stHorizontalBlock"] { align-items: center; }
            </style>
        """, unsafe_allow_html=True)
        
        # Flex-like layout using columns  
        col_nav_left, col_nav_spacer, col_nav_right = st.columns([4, 6, 2])
        
        with col_nav_left:
            # Display SVG logo in navigation
            st.markdown(f'<div style="display: flex; align-items: center; width: fit-content;">{get_logo_svg(size_px=56, font_size_px=26)}</div>', unsafe_allow_html=True)
            









            
        with col_nav_right:
            # Profile Menu logic
            # Fetch display name and user info
            user_id = st.session_state.get("user_id")
            current_user = st.session_state.get("current_user", "Account")
            display_name = st.session_state.get("user_display_name")
            
            # Normalize display name
            user_display = display_name or (current_user if isinstance(current_user, str) else current_user.get("username", "Account"))
            
            with st.popover(f"👤 {user_display}", use_container_width=True):
                st.markdown("### 👤 User Account")
                if isinstance(current_user, dict):
                    st.markdown(f"**Email:** {current_user.get('email', 'N/A')}")
                    st.markdown(f"**Role:** {current_user.get('role', 'User')}")
                else:
                    st.markdown(f"**User:** {current_user}")
                
                if st.button("👤 View Profile", key="nav_view_profile_btn", use_container_width=True):
                    st.session_state.selected_feature = "profile"
                    st.rerun()
                
                st.markdown("---")
                if st.button("🚪 Log Out", key="nav_logout_btn", use_container_width=True, type="primary"):
                    st.session_state.authenticated = False
                    st.session_state.current_user = None
                    st.session_state.user_id = None
                    st.session_state.selected_feature = None
                    st.rerun()

    st.markdown('<hr class="gradient" style="margin-top: 0; margin-bottom: 0.9rem; max-width: 100%;">', unsafe_allow_html=True)
    
    # Create Layout - use st.sidebar for proper sidebar integration
    # Features in Sidebar
    with st.sidebar:
        # st.markdown("## 🎓 Main Menu") removed in favor of expander title
        
        # Determine if expander should be open (open if no feature selected)
        is_menu_expanded = st.session_state.selected_feature is None
        
        with st.expander("Menu", expanded=True):
            # Navigation menu
            if st.button("Dashboard Home", key="nav_dashboard", use_container_width=True):
                st.session_state.selected_feature = None
                st.rerun()
            
            if st.button("Chatbot & Summarization", key="nav_chatbot", use_container_width=True):
                st.session_state.selected_feature = "chatbot"
                st.rerun()
            
            if st.button("Quiz Generator", key="nav_quiz", use_container_width=True):
                st.session_state.selected_feature = "quiz"
                st.rerun()
            
            if st.button("Live Lecture Recording", key="nav_recording", use_container_width=True):
                st.session_state.selected_feature = "recording"
                st.rerun()
            
            if st.button("Flash Cards", key="nav_flashcards", use_container_width=True):
                st.session_state.selected_feature = "flashcards"
                st.rerun()
            
            if st.button("Translation", key="nav_translation", use_container_width=True):
                st.session_state.selected_feature = "translation"
                st.rerun()
            
            if st.button("Search", key="nav_search", use_container_width=True):
                st.session_state.selected_feature = "search"
                st.rerun()
            
            if st.button("Offline Mode", key="nav_offline", use_container_width=True):
                st.session_state.selected_feature = "offline"
                st.rerun()
            
            # OTHER Section
            st.markdown('<hr style="margin: 15px 0 10px 0; border: none; height: 1px; background: #f0f4ff;">', unsafe_allow_html=True)
            st.markdown('<p class="sidebar-title" style="margin-left: 10px;">OTHER</p>', unsafe_allow_html=True)
            
            if st.button("Profile", key="nav_profile", use_container_width=True):
                st.session_state.selected_feature = "profile"
                st.rerun()

            # Custom styling for logout button
            st.markdown("""
                <style>
                /* Logout — quiet danger row */
                .st-key-logout_btn button { color: #b42318 !important; }
                .st-key-logout_btn button:hover {
                    background: #fef3f2 !important;
                    color: #b42318 !important;
                }
                </style>
            """, unsafe_allow_html=True)
            
            if st.button("Logout", key="logout_btn", help="Logout from your account", use_container_width=True):
                st.session_state.authenticated = False
                st.session_state.current_user = None
                st.session_state.selected_feature = None
                st.rerun()

    # Main Content Area
    if st.session_state.selected_feature is None:
        show_welcome_screen()
    elif st.session_state.selected_feature == "chatbot":
        show_chatbot_feature()
    elif st.session_state.selected_feature == "quiz":
        show_quiz_feature()
    elif st.session_state.selected_feature == "recording":
        show_recording_feature()
    elif st.session_state.selected_feature == "flashcards":
        show_flashcards_feature()
    elif st.session_state.selected_feature == "translation":
        show_translation_feature()
    elif st.session_state.selected_feature == "notes":
        show_notes_feature()
    elif st.session_state.selected_feature == "admin":
        show_admin_feature()
    elif st.session_state.selected_feature == "search":
        show_search_feature()
    elif st.session_state.selected_feature == "offline":
        show_offline_feature()
    elif st.session_state.selected_feature == "profile":
        show_profile_feature()
    else:
        show_coming_soon_feature(st.session_state.selected_feature)

    # Footer on every dashboard page
    render_app_footer()

def show_coming_soon_feature(feature_name):
    """Placeholder for features not yet fully implemented"""
    st.markdown(
        f"""
        <div class="coming-soon">
            <div class="coming-soon-icon">🚀</div>
            <h2 class="coming-soon-title">{feature_name.replace('_', ' ').title()} Feature</h2>
            <p class="coming-soon-text">
                We're working hard to bring you the best {feature_name.replace('_', ' ')} experience. 
                This feature will be available in the next update!
            </p>
            <div style="text-align: center; margin-top: 30px;">
                <p style="color: #667eea; font-weight: 600;">Stay tuned for educational excellence!</p>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

# ==========================
# WELCOME SCREEN SECTION
# ==========================
def show_welcome_screen():
    """Display welcome screen with learning visuals"""
    # Personalized greeting (display-only)
    _current_user = st.session_state.get("current_user")
    _greet_name = st.session_state.get("user_display_name") or (
        _current_user if isinstance(_current_user, str) else "Learner"
    )

    # PAGE HEADER
    st.markdown(
        f"""
        <div class="welcome-container">
            <p class="welcome-eyebrow">Dashboard</p>
            <h1 class="welcome-title">Welcome back, {_greet_name}</h1>
            <p class="welcome-text">
                Your comprehensive educational platform powered by AI. Choose a feature from the sidebar to get started.
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    # Learning Statistics Cards
    # Get user stats from session state (loaded during login)
    user_stats = st.session_state.get("user_stats", {})
    study_sessions = user_stats.get("study_sessions", 0)
    quizzes_created = user_stats.get("quizzes_created", 0)
    recordings_count = user_stats.get("recordings_count", 0)
    
    # Calculate progress percentage (based on activity)
    total_activities = study_sessions + quizzes_created + recordings_count
    progress_percentage = min(100, total_activities * 5) if total_activities > 0 else 0

    # Stat cards (styled by the global design system)
    st.markdown(f"""
        <div class="stats-grid">
            <div class="stat-card">
                <p class="stat-label">Study Sessions</p>
                <div class="stat-value">{study_sessions}</div>
                <p class="stat-sub">Track your learning progress</p>
            </div>
            <div class="stat-card">
                <p class="stat-label">Quizzes Created</p>
                <div class="stat-value">{quizzes_created}</div>
                <p class="stat-sub">Interactive learning materials</p>
            </div>
            <div class="stat-card">
                <p class="stat-label">Recordings</p>
                <div class="stat-value">{recordings_count}</div>
                <p class="stat-sub">Audio content processed</p>
            </div>
            <div class="stat-card">
                <p class="stat-label">Progress</p>
                <div class="stat-value">{progress_percentage}%</div>
                <p class="stat-sub">Learning efficiency</p>
            </div>
        </div>
    """, unsafe_allow_html=True)

    # Inspirational Quote (Themed BG)
    quote_data = get_random_quote()
    st.markdown(
        f"""
        <div class="quote-card">
            <h3>Today's Learning Quote</h3>
            <p class="quote-text">"{quote_data['text']}"</p>
            <p class="quote-author">— {quote_data['author']}</p>
        </div>
        """,
        unsafe_allow_html=True
    )

# ==========================
# CHATBOT AND SUMMARIZATION SECTION
# ==========================
def show_chatbot_feature():

    
    
    # ✅ Configurable Tesseract OCR path with fallbacks
    # Priority: ENV var TESSERACT_CMD -> Windows default path -> system PATH
    _tesseract_env = os.getenv("TESSERACT_CMD")
    _windows_default = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
    try:
        if _tesseract_env and os.path.exists(_tesseract_env):
            pytesseract.pytesseract.tesseract_cmd = _tesseract_env
        elif os.name == "nt" and os.path.exists(_windows_default):
            pytesseract.pytesseract.tesseract_cmd = _windows_default
        # else: rely on PATH; if not present, OCR calls will raise which we handle
    except Exception:
        # If configuration fails, we let _process_image handle exceptions gracefully
        pass
    
    
    class DocumentProcessor:
        def __init__(self):
            self.supported_formats = {
                ".txt", ".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg"
            }
    
        def process_document(self, filepath):
            """
            Detect file type and extract text safely.
            Always returns a string (even if it's an error).
            """
            try:
                file_extension = os.path.splitext(filepath)[1].lower()
    
                if file_extension == ".txt":
                    return self._process_txt(filepath)
                elif file_extension == ".pdf":
                    return self._process_pdf(filepath)
                elif file_extension in [".docx", ".doc"]:
                    return self._process_word(filepath)
                elif file_extension in [".png", ".jpg", ".jpeg"]:
                    return self._process_image(filepath)
                else:
                    return f"[Unsupported file format: {file_extension}]"
    
            except Exception as e:
                return f"[File processing error: {str(e)}]"
    
        # ------------------------
        # File Type Processors
        # ------------------------
    
        def _process_txt(self, filepath):
            """Extract text from .txt files."""
            try:
                with open(filepath, "r", encoding="utf-8") as file:
                    content = file.read()
                return self._clean_text(content)
            except UnicodeDecodeError:
                try:
                    with open(filepath, "r", encoding="latin-1") as file:
                        content = file.read()
                    return self._clean_text(content)
                except Exception as e:
                    return f"[Error reading TXT file: {str(e)}]"
    
        def _process_pdf(self, filepath):
            """Extract text from PDF files."""
            try:
                with open(filepath, "rb") as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    content = ""
                    for page in pdf_reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            content += extracted + "\n"
                return self._clean_text(content) if content else "[No text extracted from PDF]"
            except Exception as e:
                return f"[Error reading PDF: {str(e)}]"
    
        def _process_word(self, filepath):
            """Extract text from Word documents (.docx and .doc)."""
            try:
                doc = Document(filepath)
                content = ""
    
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
    
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            content += cell.text + " "
                    content += "\n"
    
                return self._clean_text(content) if content else "[No text extracted from Word file]"
            except Exception as e:
                return f"[Error reading Word document: {str(e)}]"
    
        def _process_image(self, filepath):
            """Extract text from image files using OCR (safe with timeout)."""
            try:
                with Image.open(filepath) as img:
                    img = img.convert("RGB")  # ensure format is consistent
                    # Add timeout to prevent hanging on large images
                    content = pytesseract.image_to_string(img, timeout=30)
                    cleaned = self._clean_text(content)
                    return cleaned if cleaned else "[No text detected in image]"
            except Exception as e:
                # Be explicit when Tesseract is missing to help users
                if "tesseract is not installed" in str(e).lower() or "not found" in str(e).lower():
                    return "[OCR unavailable: Tesseract not found. Install Tesseract or set TESSERACT_CMD]"
                return f"[Error reading image: {str(e)}]"
    
        # ------------------------
        # Helpers
        # ------------------------
    
        def _clean_text(self, text):
            """Clean and normalize extracted text."""
            if not text:
                return ""
    
            # Remove extra whitespace
            text = re.sub(r"\s+", " ", text)
    
            # Remove unwanted characters but keep punctuation
            text = re.sub(r"[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}]", "", text)
    
            # Normalize spaces
            text = re.sub(r"\s+", " ", text)
    
            return text.strip()
    
        def get_document_summary(self, content, max_length=500):
            """Generate a brief summary of the document content."""
            if not content:
                return "[No content available to summarize]"
    
            if len(content) <= max_length:
                return content
    
            # Take first few sentences
            sentences = content.split(".")
            summary = ""
            for sentence in sentences:
                if len(summary + sentence) < max_length:
                    summary += sentence.strip() + ". "
                else:
                    break
    
            return summary.strip()
    
    
    
    
    
    
    doc_processor = DocumentProcessor()
    
    
    
    # ---------------------------
    # Page Config
    # ---------------------------
    # st.set_page_config(page_title="LectureBuddies - AI Chatbot", page_icon="🤖", layout="wide", initial_sidebar_state="expanded")
    
    # ---------------------------
    # Load API Key from .env
    # ---------------------------
    load_dotenv()
    api_key = os.getenv("GROQ_API_KEY")  # make sure your .env has GROQ_API_KEY=your_api_key_here
    
    if not api_key:
        st.error("⚠️ API key missing! Please check your .env file.")
        st.stop()
    
    # ---------------------------
    # Session Initialization
    # ---------------------------
    def init_session_state():
        defaults = {
            "messages": [],
            "uploaded_files": [],
            "document_contents": {}
        }
        for k, v in defaults.items():
            if k not in st.session_state:
                st.session_state[k] = v
    
    init_session_state()
    
    # ---------------------------
    # Inject File Content Helper
    # ---------------------------
    def inject_file_content(user_message: str) -> str:
        """
        Replace file references in user message with extracted text,
        so model never says 'I can't see images'.
        """
        for fname, content in st.session_state.document_contents.items():
            if fname.lower() in user_message.lower():
                extracted = content if content.strip() else "[No text extracted from this file]"
                user_message = user_message.replace(
                    fname,
                    f"(Extracted content: {extracted[:1000]}...)"
                )
        return user_message
    
    # ---------------------------
    # API Interaction
    # ---------------------------
    def get_groq_response(user_input, model="llama-3.1-8b-instant"):
    
        """Send query + context to Groq API and return assistant response."""
        if not api_key or api_key.strip() == "":
            return "⚠️ Missing API key. Please set GROQ_API_KEY in your .env file."
    
        url = "https://api.groq.com/openai/v1/chat/completions"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    
        # Build document context
        doc_context = ""
        if st.session_state.document_contents:
            doc_context = "\n\n**Available Documents:**\n"
            for fname, content in st.session_state.document_contents.items():
                snippet = content[:1000] if content else "[No extractable text]"
                doc_context += f"\n--- {fname} ---\n{snippet}...\n"
    
        # Build conversation
        messages = [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages]
        system_msg = (
            f"You are Lecturebuddies, an AI chatbot designed for education. "
            f"Answer clearly, summarize effectively, and explain concepts step by step."
            f"{' Available Documents: ' + doc_context if doc_context else ''}\n\n"
            "Guidelines:\n"
            "📚 Education-focused\n"
            "📝 Summarization expert\n"
            "🎯 Clarity first (simple language, then details)\n"
            "✅ Confidence + accuracy\n"
            "Break down topics step-by-step, use examples, and stay professional yet supportive."
        )
        messages.insert(0, {"role": "system", "content": system_msg})
        messages.append({"role": "user", "content": user_input})
    
        payload = {"model": model, "messages": messages, "temperature": 0.7, "max_tokens": 1000}
    
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=30)
            if resp.status_code == 200:
                data = resp.json()
                return data.get("choices", [{}])[0].get("message", {}).get("content", "⚠️ No response received.")
            elif resp.status_code == 401:
                return "❌ Invalid API key. Please check GROQ_API_KEY in your .env file."
            elif resp.status_code == 429:
                return "⏳ Too many requests. Please slow down and retry shortly."
            else:
                return f"⚠️ API Error {resp.status_code}: {resp.text}"
        except requests.exceptions.Timeout:
            return "⏳ Request timed out. Please retry."
        except requests.exceptions.RequestException as e:
            return f"🌐 Network error: {e}"
        except Exception as e:
            return f"⚠️ Unexpected error: {e}"
    
    # ---------------------------
    # Document Processing
    # ---------------------------
    def process_document(uploaded_file):
        """Extract text from uploaded documents (txt, pdf, docx, images with OCR)."""
        try:
            # Save the uploaded file temporarily
            temp_path = os.path.join("temp", uploaded_file.name)
            os.makedirs("temp", exist_ok=True)
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.read())
    
            # Use your DocumentProcessor
            content = doc_processor.process_document(temp_path)
    
            return content if content.strip() else "[No text extracted]"
        except Exception as e:
            return f"[File processing error: {e}]"
    
    # ---------------------------
    # Enhanced Styling (Matching Quiz Generator Theme)
    # ---------------------------
    st.markdown(
        """
        <style>
        /* Chat page: bubbles, header and inputs are styled by the global
           design system. Page-specific tweaks only. */
        .chat-header {
            color: #16182d;
            font-size: 20px;
            font-weight: 700;
            margin-bottom: 10px;
            text-align: center;
        }
        </style>
        """,
        unsafe_allow_html=True
    )
    
    # ---------------------------
    # Sidebar (Enhanced to Match Theme)
    # ---------------------------
    with st.sidebar:
        st.markdown("## ⚙️ Chat Settings")
        st.markdown("### Customize your experience")
        
        if st.button("🗑️ Clear Chat", key="clear_chat", help="Start a new conversation"):
            st.session_state.messages.clear()
            st.rerun()
    
        st.markdown("---")
        st.markdown("**📎 Upload Files**")
        sidebar_upload = st.file_uploader(
            "Choose files",
            type=['txt', 'pdf', 'docx', 'doc', 'png', 'jpg', 'jpeg', 'gif', 'bmp'],
            key="sidebar_uploader",
            label_visibility="collapsed",
            help="Upload documents or images for context (PDF, DOCX, TXT, Images with OCR)"
        )
        if sidebar_upload and sidebar_upload.name not in st.session_state.document_contents:
            file_details = {
                "filename": sidebar_upload.name,
                "filetype": sidebar_upload.type,
                "filesize": sidebar_upload.size
            }
            with st.spinner(f"Processing {sidebar_upload.name}..."):
                st.session_state.document_contents[sidebar_upload.name] = process_document(sidebar_upload)
            st.session_state.uploaded_files.append(file_details)
            st.sidebar.success(f"✅ {sidebar_upload.name} uploaded!")
            st.rerun()
    
        # Sidebar tips (Compact)
        st.markdown("---")
        st.markdown("**💡 Quick Tips:**")
        st.markdown("- Ask about studies or homework")
        st.markdown("- Upload files for context")
        st.markdown("- Be specific for better responses")
    
        # Show uploaded files (Styled)
        if st.session_state.uploaded_files:
            st.markdown("---")
            st.markdown("**📁 Your Files:**")
            for i, f in enumerate(st.session_state.uploaded_files):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"📄 {f['filename']}")
                with col2:
                    if st.button("🗑️", key=f"del_file_{i}", help="Remove file"):
                        fname = f['filename']
                        st.session_state.uploaded_files.pop(i)
                        st.session_state.document_contents.pop(fname, None)
                        st.rerun()
    

    
   
    
    # ---------------------------
    # Quick Actions (Styled to Match)
    # ---------------------------
    if not st.session_state.messages:
        st.markdown(
            """
            <div class="welcome-container">
                <h1 class="welcome-title">Lecturebuddies Chatbot</h1>
                <p class="welcome-text">Start chatting or try a quick action below to dive into your studies.</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        
        col1, col2, col3 = st.columns(3)
        presets = {
            "📚 Help with Homework": "I need help understanding this homework assignment. Can you explain step-by-step?",
            "🔬 Explain a Concept": "I'm studying this concept but finding it difficult. Can you explain clearly with examples?",
            "💡 Study Tips": "I want to improve my study efficiency. What study strategies should I use?"
        }
        for col, (label, prompt) in zip([col1, col2, col3], presets.items()):
            with col:
                if st.button(label, key=label.replace(" ", "_").lower(), help="Start with this prompt"):
                    st.session_state.messages.append({"role": "user", "content": prompt})
                    with st.spinner("🤔 Thinking..."):
                        reply = get_groq_response(prompt)
                    st.session_state.messages.append({"role": "assistant", "content": reply})
                    st.rerun()
    
    # ---------------------------
    # Chat Display (Direct, No Container)
    # ---------------------------
    if st.session_state.messages:
        for msg in st.session_state.messages:
            if msg["role"] == "user":
                st.markdown(f"""
                <div class="user-message">
                    <div class="user-bubble">
                        {msg['content']}
                    </div>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div class="assistant-message">
                    <div class="assistant-bubble">
                        {msg['content']}
                    </div>
                </div>
                """, unsafe_allow_html=True)
    
    # ---------------------------
    # Chat Input (Professional Styling, Shorter Placeholder)
    # ---------------------------
    user_input = st.chat_input(placeholder="💬 Ask about studies or uploaded files...")
    if user_input and user_input.strip():
        # 🔑 Inject file content here
        final_input = inject_file_content(user_input.strip())
    
        st.session_state.messages.append({"role": "user", "content": user_input.strip()})
        with st.spinner("🤔 Thinking..."):
            reply = get_groq_response(final_input)
        st.session_state.messages.append({"role": "assistant", "content": reply})
        st.rerun()
    
    # Footer is rendered globally by render_app_footer()

# ==========================
# QUIZ GENERATOR SECTION
# ==========================
def show_quiz_feature():
    """Display quiz generator feature with old design"""
    # Layout: Settings in Sidebar (like Chatbot), Main Content in Container
    # col_feature_sidebar, col_main_content = st.columns([1, 3]) # Removed to use genuine sidebar
    
    # Remove bottom spacing for this view
    st.markdown(
        """
        <style>
        section[data-testid="stMain"] .block-container { padding-bottom: 0 !important; margin-bottom: 0 !important; }
        </style>
        """,
        unsafe_allow_html=True
    )
    
    # LEFT: Feature-specific sidebar (Quiz Settings)
    with st.sidebar:
        st.markdown("### ⚙️ Quiz Settings")
        st.markdown("Customize your quiz generation.")
        
        st.markdown("**📊 Number of Questions**")
        num_questions = st.slider("How many questions?", min_value=1, max_value=20, value=st.session_state.num_questions, key="num_slider", label_visibility="collapsed")
        st.session_state.num_questions = num_questions

        st.markdown("**🎯 Difficulty Level**")
        difficulty = st.selectbox("Select difficulty:", ["Easy", "Medium", "Hard"], index=["Easy", "Medium", "Hard"].index(st.session_state.difficulty), key="diff_select", label_visibility="collapsed")
        st.session_state.difficulty = difficulty

        # Model and temperature controls
        model_options = [
            "llama-3.1-8b-instant",
            "llama-3.1-70b-versatile",
            "llama-3.2-11b-text-preview"
        ]
        st.session_state.quiz_model = st.selectbox("Model", model_options, index=model_options.index(st.session_state.quiz_model) if st.session_state.quiz_model in model_options else 0)
        st.session_state.quiz_temperature = st.slider("Creativity (temperature)", 0.0, 1.0, float(st.session_state.quiz_temperature), 0.1)
    
    # RIGHT: Main Quiz Interface
    with st.container():
        # Header
        st.markdown(
            """
            <div class="welcome-container">
                <h1 class="welcome-title">Quiz Generator</h1>
                <p class="welcome-text">
                    Transform your notes, lectures, or ideas into interactive quizzes with AI magic ✨
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )
        
        tab1, tab2, tab3 = st.tabs(["📁 Upload File", "💡 Enter Prompt", "📋 Paste Text"])

        # Tab 1 - File Upload
        with tab1:
            st.markdown("### Upload a document to generate a quiz from its content")
            uploaded_file = st.file_uploader(
                "Choose a file (PDF, DOCX, TXT)",
                type=["pdf", "docx", "txt"],
                help="Supported formats: PDF, Word documents, and plain text files."
            )

            if st.button("✨ Generate Quiz from File", key="file-btn", help=f"Create {num_questions} {difficulty} questions"):
                if uploaded_file:
                    content = extract_content_from_file(uploaded_file)
                    if content and "Error" not in content:
                        quiz_result = get_groq_quiz_response(content, num_questions, difficulty, model=st.session_state.quiz_model, temperature=st.session_state.quiz_temperature)
                        st.session_state.quiz_output = quiz_result
                        # Save quiz to database
                        if st.session_state.get("user_id") and "Error" not in quiz_result:
                            save_quiz(st.session_state.user_id, uploaded_file.name, num_questions, difficulty, quiz_result)
                            # Refresh user stats
                            st.session_state.user_stats = get_user_stats(st.session_state.user_id)
                    else:
                        st.error("Failed to extract content from file. Please try a different file or format.")
                else:
                    st.warning("Please upload a file first!")

        # Tab 2 - Prompt
        with tab2:
            st.markdown("### Describe a topic or provide content for the quiz")
            prompt_text = st.text_area(
                "Enter your topic, subject, or detailed content",
                placeholder="e.g., 'Explain photosynthesis and generate questions on it' or paste lecture notes...",
                height=100,
                help="Be as detailed as possible for better quizzes!"
            )

            if st.button("✨ Generate Quiz from Prompt", key="prompt-btn", help=f"Create {num_questions} {difficulty} questions"):
                if prompt_text.strip():
                    quiz_result = get_groq_quiz_response(prompt_text, num_questions, difficulty, model=st.session_state.quiz_model, temperature=st.session_state.quiz_temperature)
                    st.session_state.quiz_output = quiz_result
                    # Save quiz to database
                    if st.session_state.get("user_id") and not quiz_result.startswith("Error"):
                        save_quiz(st.session_state.user_id, f"Prompt: {prompt_text[:50]}...", num_questions, difficulty, quiz_result)
                        # Save to local for offline access
                        local_data = {
                            "subject": f"Prompt: {prompt_text[:20]}...",
                            "num_questions": num_questions,
                            "difficulty": difficulty,
                            "content": quiz_result
                        }
                        save_to_local(st.session_state.current_user, "quizzes", local_data)
                        
                        # Refresh user stats
                        st.session_state.user_stats = get_user_stats(st.session_state.user_id)
                else:
                    st.warning("Please enter some content or a topic!")

        # Tab 3 - Text Input
        with tab3:
            st.markdown("### Paste scanned or copied text directly")
            scanned_text = st.text_area(
                "Paste your text content here",
                placeholder="e.g., Copy-paste from a scanned PDF, image OCR, or notes...",
                height=100,
                help="Ideal for quick text from any source."
            )

            if st.button("✨ Generate Quiz from Text", key="scan-btn", help=f"Create {num_questions} {difficulty} questions"):
                if scanned_text.strip():
                    quiz_result = get_groq_quiz_response(scanned_text, num_questions, difficulty, model=st.session_state.quiz_model, temperature=st.session_state.quiz_temperature)
                    st.session_state.quiz_output = quiz_result
                    # Save quiz to database
                    if st.session_state.get("user_id") and not quiz_result.startswith("Error"):
                        save_quiz(st.session_state.user_id, "Text Input", num_questions, difficulty, quiz_result)
                        # Save to local for offline access
                        local_data = {
                            "subject": "Text Input Quiz",
                            "num_questions": num_questions,
                            "difficulty": difficulty,
                            "content": quiz_result
                        }
                        save_to_local(st.session_state.current_user, "quizzes", local_data)
                        
                        # Refresh user stats
                        st.session_state.user_stats = get_user_stats(st.session_state.user_id)
                else:
                    st.warning("Please paste some text!")

        # Display Quiz Results
        if st.session_state.quiz_output:
            if "Error" in st.session_state.quiz_output or "⚠️" in st.session_state.quiz_output or "❌" in st.session_state.quiz_output:
                st.error(st.session_state.quiz_output)
                if st.button("Clear", key="clear_error", help="Start over"):
                    st.session_state.quiz_output = None
                    st.rerun()
            else:
                st.markdown(
                    """
                    <div class="quiz-container">
                        <h3 style="color: #4e54c8; font-size: 22px; font-weight: 700; margin-bottom: 10px; text-align: center; font-family: 'Georgia', serif;">
                            Your Generated Quiz
                        </h3>
                        <p style="text-align: center; color: #666; font-style: italic; font-size: 14px;">
                            {num_questions} questions at {difficulty} difficulty level
                        </p>
                    </div>
                    """.format(num_questions=st.session_state.num_questions, difficulty=st.session_state.difficulty),
                    unsafe_allow_html=True
                )

                st.markdown(st.session_state.quiz_output)

                st.download_button(
                    label="Download Quiz as TXT",
                    data=st.session_state.quiz_output,
                    file_name=f"lecturebuddies_quiz_{st.session_state.num_questions}q_{st.session_state.difficulty.lower()}.txt",
                    mime="text/plain",
                    help="Save your quiz for later use!"
                )

                if st.button("Generate New Quiz", key="clear", help="Clear and start over"):
                    st.session_state.quiz_output = None
                    st.rerun()

# ==========================
# LIVE LECTURE RECORDING SECTION
# ==========================
def show_recording_feature():
    """Display simplified speech-to-text transcription via file upload or Real-Time"""
    
    # 1. Main heading
    st.markdown(
        """
        <div class="welcome-container">
            <h1 class="welcome-title">Speech to Text</h1>
            <p class="welcome-text">
                Upload voice recordings or record live audio for real-time transcription
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    SAVE_DIR = "temp_recordings"
    os.makedirs(SAVE_DIR, exist_ok=True)

    # 2. Load Whisper Model (Updated to 'tiny.en' for speed as requested)
    @st.cache_resource
    def load_whisper_model_simple(model_size="tiny.en", device="cpu", compute_type="int8"):
        # Try to use GPU if available, else CPU
        try:
            return WhisperModel(model_size, device="cuda", compute_type="int8_float16")
        except:
            return WhisperModel(model_size, device="cpu", compute_type="int8")

    with st.spinner(f"Loading High-Speed Model..."):
        model = load_whisper_model_simple()

    # 3. Tabs for different modes
    tab1, tab2 = st.tabs(["📁 Upload Audio Files", "🎤 Realtime Transcript"])

    # ==========================
    # TAB 1: Upload Audio Files (Existing Working Code)
    # ==========================
    with tab1:
        st.markdown("### Upload audio files for transcription")
        
        uploaded_files = st.file_uploader(
            "Upload audio file(s)",
            type=["wav", "mp3", "m4a", "flac", "ogg"],
            help="Supported formats: WAV, MP3, M4A, FLAC, OGG",
            accept_multiple_files=True,
            key="upload_files"
        )

        if uploaded_files:
            if st.button("🎯 Transcribe Files", key="transcribe_all"):
                for idx, uploaded_file in enumerate(uploaded_files, start=1):
                    with st.spinner(f"Transcribing {uploaded_file.name} ({idx}/{len(uploaded_files)})..."):
                        # Save to temp
                        temp_path = os.path.join(SAVE_DIR, f"uploaded_{int(time.time())}_{uploaded_file.name}")
                        with open(temp_path, "wb") as f:
                            f.write(uploaded_file.read())

                        try:
                            segments, info = model.transcribe(temp_path, beam_size=5)
                            transcription_text = "".join([seg.text for seg in segments])

                            st.success(f"✅ Transcription completed: {uploaded_file.name}")

                            # Display and downloads
                            with st.expander(f"📝 {uploaded_file.name} — View Transcription", expanded=True):
                                st.text_area("Result", value=transcription_text, height=200, key=f"tx_{idx}")
                                st.download_button(
                                    "📥 Download Transcription",
                                    transcription_text,
                                    file_name=f"transcript_{uploaded_file.name}.txt"
                                )
                        except Exception as e:
                            st.error(f"❌ Error transcribing {uploaded_file.name}: {str(e)}")
                        finally:
                            if os.path.exists(temp_path):
                                os.remove(temp_path)

    # ==========================
    # TAB 2: Realtime Transcript (UPDATED CODE)
    # ==========================
    # with tab2:
    #     st.markdown("### ⚡ Fast Real-time Transcription")
        
    #     # Initialize session state for transcript storage
    #     if "live_transcript" not in st.session_state:
    #         st.session_state.live_transcript = ""

    # ==========================
    # TAB 2: Realtime Transcript (UPDATED CODE)
    # ==========================
    with tab2:
        st.markdown("### ⚡ Fast Real-time Transcription")
        
        # Initialize session state for transcript storage
        if "live_transcript" not in st.session_state:
            st.session_state.live_transcript = ""
        if "is_recording" not in st.session_state:
            st.session_state.is_recording = False

        col1, col2 = st.columns([1, 2])

        with col1:
            st.info("""💡 **Online Live Recording:**
1. Click **'Start Recording'** below.
2. **Allow microphone access** in your browser.
3. Speak clearly into your microphone.
4. Click **'Stop'** to instantly transcribe the audio.""")
            
            # Browser-based microphone recorder (Works on Streamlit Cloud!)
            audio = mic_recorder(
                start_prompt="🔴 Start Recording",
                stop_prompt="⏹️ Stop Recording",
                just_once=False,
                use_container_width=True,
                key='browser_recorder'
            )
            
            if audio:
                audio_bytes = audio['bytes']
                if audio_bytes:
                    with st.spinner("⚡ Transcribing your recording..."):
                        # Use a temporary file to process the audio
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
                            tmp_file.write(audio_bytes)
                            tmp_path = tmp_file.name
                        
                        try:
                            # Process with Faster-Whisper
                            segments, _ = model.transcribe(tmp_path, beam_size=1, language="en")
                            new_text = "".join([s.text for s in segments]).strip()
                            
                            if new_text:
                                st.session_state.live_transcript += new_text + " "
                                
                                # Automatic Session Saving
                                session_start_time = int(time.time())
                                session_filename = f"recording_{session_start_time}_transcripts.json"
                                local_data = {
                                    "title": f"Live Recording {time.strftime('%H:%M')}",
                                    "content": st.session_state.live_transcript,
                                    "saved_at": session_start_time
                                }
                                save_to_local(st.session_state.current_user, "transcripts", local_data, custom_filename=session_filename)
                                st.success("✅ Added to Transcript!")
                        except Exception as e:
                            st.error(f"Transcription error: {e}")
                        finally:
                            if os.path.exists(tmp_path):
                                os.remove(tmp_path)
            
            st.markdown("---")
            
            # Download Button (Visible if we have text)
            if st.session_state.live_transcript:
                st.download_button(
                    label="📥 Download Transcript",
                    data=st.session_state.live_transcript,
                    file_name="live_transcript.txt",
                    mime="text/plain",
                    use_container_width=True
                )
                
            if st.button("🗑️ Clear Transcript", use_container_width=True):
                st.session_state.live_transcript = ""
                st.rerun()

        with col2:
            st.markdown("#### 📝 Live Output")
            # Create a placeholder to show the transcript
            st.text_area(
                "Transcript", 
                value=st.session_state.live_transcript, 
                height=450, 
                placeholder="Transcribed text will appear here after you stop recording..."
            )






# ==========================
# FLASH CARDS FEATURE
# ==========================
def show_flashcards_feature():
    """Display flash cards generator feature"""
    st.markdown(
        """
        <div class="welcome-container">
            <h1 class="welcome-title">Flash Cards Generator</h1>
            <p class="welcome-text">
                Create interactive flash cards from your study materials
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### ⚙️ Settings")
        st.markdown("---")
        
        # Flash card settings
        num_cards = st.slider("Number of Cards", 1, 50, 10)
        difficulty = st.selectbox("Difficulty", ["Easy", "Medium", "Hard"])
        subject = st.text_input("Subject/Topic", placeholder="e.g., Biology, History")
        
        st.markdown("---")
        st.markdown("**📄 Upload Content**")
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['txt', 'pdf', 'docx'],
            help="Upload study material to generate flash cards"
        )
    
    with col2:
        st.markdown("### 📝 Content Input")
        
        content_input = st.text_area(
            "Enter or paste content for flashcards",
            placeholder="Paste your study material here or upload a file...",
            height=200,
            help="Provide content to generate flashcards from"
        )
        
        if st.button("🎯 Generate Flash Cards", key="generate_cards", use_container_width=True):
            if content_input or uploaded_file:
                with st.spinner("Generating flash cards..."):
                    # Extract content from file if uploaded, otherwise use text input
                    if uploaded_file:
                        content = extract_content_from_file(uploaded_file)
                        if content and "Error" not in content:
                            flashcards = generate_flashcards(content, num_cards, difficulty, subject)
                        else:
                            st.error("Failed to extract content from file. Please try a different file.")
                            flashcards = None
                    else:
                        flashcards = generate_flashcards(content_input, num_cards, difficulty, subject)
                    
                    # Only save if flashcards were successfully generated
                    if flashcards:
                        st.session_state.flashcards = flashcards
                        
                        # Save to local for offline access
                        local_data = {
                            "subject": subject or "General Flashcards",
                            "difficulty": difficulty,
                            "cards": flashcards
                        }
                        save_to_local(st.session_state.current_user, "flashcards", local_data)
                        
                        st.success(f"✅ Generated {len(flashcards)} flash cards!")
            else:
                st.warning("Please provide content to generate flash cards")
        
        # Display generated flash cards
        if st.session_state.get('flashcards'):
            st.markdown("### 🃏 Your Flash Cards")
            display_flashcards(st.session_state.flashcards)

def generate_flashcards(content, num_cards, difficulty, subject):
    """Generate flash cards from content using AI"""
    if not api_key:
        return [{"front": "Error: No API key", "back": "Please set GROQ_API_KEY in your .env file"}]
    
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    
    prompt = f"""
    Create {num_cards} flash cards from this content for {subject or 'general study'}.
    Difficulty level: {difficulty}
    
    Content: {content[:2000]}
    
    Format each card as:
    Front: [Question or term]
    Back: [Answer or definition]
    
    Make them educational and useful for studying.
    """
    
    messages = [
        {"role": "system", "content": "You are a helpful study assistant that creates educational flash cards."},
        {"role": "user", "content": prompt}
    ]
    
    payload = {"model": "llama-3.1-8b-instant", "messages": messages, "temperature": 0.7, "max_tokens": 1000}
    
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=30)
        if response.status_code == 200:
            data = response.json()
            content = data.get("choices", [{}])[0].get("message", {}).get("content", "")
            return parse_flashcards(content)
        else:
            return [{"front": "Error generating cards", "back": f"API Error: {response.status_code}"}]
    except Exception as e:
        return [{"front": "Error", "back": f"Failed to generate cards: {str(e)}"}]

def parse_flashcards(content):
    """Parse AI response into flash card format"""
    cards = []
    lines = content.split('\n')
    current_card = {}
    
    for line in lines:
        line = line.strip()
        if line.startswith('Front:'):
            if current_card:
                cards.append(current_card)
            current_card = {"front": line.replace('Front:', '').strip(), "back": ""}
        elif line.startswith('Back:'):
            current_card["back"] = line.replace('Back:', '').strip()
    
    if current_card:
        cards.append(current_card)
    
    return cards if cards else [{"front": "Sample Question", "back": "Sample Answer"}]

def display_flashcards(cards):
    """Display flash cards in an interactive format"""
    if not cards:
        return
    
    # Initialize session state for current card
    if 'current_card_index' not in st.session_state:
        st.session_state.current_card_index = 0
    
    current_index = st.session_state.current_card_index
    current_card = cards[current_index]
    
    # Card display
    st.markdown(
        f"""
        <div class="flash-card">
            <div class="flash-card-counter">Card {current_index + 1} of {len(cards)}</div>
            <div class="flash-card-face">
                <h4>Front</h4>
                <p>{current_card['front']}</p>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Show answer button
    if st.button("👁️ Show Answer", key="show_answer"):
        st.markdown(
            f"""
            <div class="flash-card-face answer" style="margin: 16px 0;">
                <h4>Back</h4>
                <p>{current_card['back']}</p>
            </div>
            """,
            unsafe_allow_html=True
        )
    
    # Navigation
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col1:
        if st.button("⬅️ Previous", key="prev_card", disabled=current_index == 0):
            st.session_state.current_card_index = max(0, current_index - 1)
            st.rerun()
    
    with col2:
        st.markdown(f"<div style='text-align: center; color: #666;'>Card {current_index + 1} of {len(cards)}</div>", unsafe_allow_html=True)
    
    with col3:
        if st.button("Next ➡️", key="next_card", disabled=current_index == len(cards) - 1):
            st.session_state.current_card_index = min(len(cards) - 1, current_index + 1)
            st.rerun()

# ==========================
# MULTILINGUAL TRANSLATION FEATURE
# ==========================
def show_translation_feature():
    """Display multilingual translation feature"""
    st.markdown(
        """
        <div class="welcome-container">
            <h1 class="welcome-title"> Multilingual Translation</h1>
            <p class="welcome-text">
                Translate text to your selected language (source auto-detected)
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### ⚙️ Translation Settings")
        st.markdown("---")
        
        # Language selection (target only)
        languages = {
            "English": "en",
            "Spanish": "es", 
            "French": "fr",
            "German": "de",
            "Italian": "it",
            "Portuguese": "pt",
            "Chinese": "zh",
            "Japanese": "ja",
            "Korean": "ko",
            "Arabic": "ar",
            "Hindi": "hi",
            "Russian": "ru",
            "Urdu": "ur",
            "Punjabi": "pa",
            "Sindhi": "sd"
        }
        target_lang_label = st.selectbox("To Language", list(languages.keys()))
        
        st.markdown("---")
        st.markdown("**📄 Upload Document**")
        uploaded_file = st.file_uploader(
            "Choose a text file",
            type=['txt', 'docx'],
            help="Upload a document to translate"
        )
    
    with col2:
        st.markdown("### 📝 Text Translation")
        
        # Text input
        text_input = st.text_area(
            "Enter text to translate",
            placeholder="Type or paste your text here...",
            height=200
        )
        
        if "translation_version" not in st.session_state:
            st.session_state.translation_version = 0

        if st.button("🔄 Translate", key="translate_text", use_container_width=True):
            if text_input or uploaded_file:
                with st.spinner("Translating..."):
                    if uploaded_file:
                        # Process uploaded file
                        content = extract_content_from_file(uploaded_file)
                        translated = translate_text(content, target_lang_label)
                    else:
                        translated = translate_text(text_input, target_lang_label)
                    
                    st.session_state.translation_result = translated
                    
                    # Increment version to force widget refresh
                    st.session_state.translation_version += 1
                        
                    st.success("✅ Translation completed!")
                    st.rerun()
            else:
                st.warning("Please provide text to translate")
        
        # Display translation result
        if st.session_state.get('translation_result'):
            st.markdown("### 🌐 Translation Result")
            # Dynamic key forces re-render
            current_key = f"translation_display_{st.session_state.translation_version}"
            st.text_area("Translated Text", value=st.session_state.translation_result, height=200, key=current_key)
            
            # Download button
            st.download_button(
                "📥 Download Translation",
                st.session_state.translation_result,
                file_name=f"translation_to_{languages[target_lang_label]}.txt",
                mime="text/plain"
            )

def translate_text(text, target_language_label):
    """Translate text using AI; auto-detect source language"""
    if not api_key:
        return "Error: No API key available"
    
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    
    prompt = (
        f"Detect the source language and translate the following text to {target_language_label}. "
        f"Only return the translated text, no explanations or prefixes.\n\n{text}"
    )
    
    messages = [
        {"role": "system", "content": "You are a professional translator. Translate accurately and naturally."},
        {"role": "user", "content": prompt}
    ]
    
    payload = {"model": "llama-3.1-8b-instant", "messages": messages, "temperature": 0.3, "max_tokens": 1000}
    
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=30)
        if response.status_code == 200:
            data = response.json()
            return data.get("choices", [{}])[0].get("message", {}).get("content", "Translation failed")
        else:
            return f"Translation error: {response.status_code}"
    except Exception as e:
        return f"Translation failed: {str(e)}"

# ==========================
# NOTES MANAGER FEATURE
# ==========================
def show_notes_feature():
    """Display notes manager with import/export and organization"""
    st.markdown(
        """
        <div class="welcome-container">
            <h1 class="welcome-title">📋 Notes Manager</h1>
            <p class="welcome-text">
                Organize, import, and export your study notes with rich text formatting
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Initialize notes in session state
    if 'notes' not in st.session_state:
        st.session_state.notes = []
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### 📁 Note Organization")
        st.markdown("---")
        
        # Categories
        categories = ["General", "Mathematics", "Science", "History", "Language", "Other"]
        selected_category = st.selectbox("Category", categories)
        
        # Tags
        tags_input = st.text_input("Tags (comma-separated)", placeholder="e.g., important, exam, review")
        
        st.markdown("---")
        st.markdown("**📤 Import/Export**")
        
        # Import notes
        uploaded_notes = st.file_uploader(
            "Import Notes",
            type=['txt', 'json'],
            help="Import notes from file"
        )
        
        if uploaded_notes:
            if st.button("📥 Import Notes", key="import_notes"):
                try:
                    if uploaded_notes.type == "application/json":
                        notes_data = json.loads(uploaded_notes.read())
                        st.session_state.notes.extend(notes_data)
                    else:
                        content = uploaded_notes.read().decode('utf-8')
                        new_note = {
                            "title": uploaded_notes.name,
                            "content": content,
                            "category": selected_category,
                            "tags": tags_input.split(',') if tags_input else [],
                            "created": time.time()
                        }
                        st.session_state.notes.append(new_note)
                    st.success("✅ Notes imported successfully!")
                except Exception as e:
                    st.error(f"❌ Import failed: {str(e)}")
        
        # Export notes
        if st.session_state.notes:
            if st.button("📤 Export All Notes", key="export_notes"):
                notes_json = json.dumps(st.session_state.notes, indent=2)
                st.download_button(
                    "💾 Download Notes (JSON)",
                    notes_json,
                    file_name=f"lecturebuddies_notes_{int(time.time())}.json",
                    mime="application/json"
                )
    
    with col2:
        st.markdown("### ✏️ Create New Note")
        
        # Note creation form
        note_title = st.text_input("Note Title", placeholder="Enter note title")
        
        # Rich text editor (simplified)
        note_content = st.text_area(
            "Note Content",
            placeholder="Write your note here...\n\nYou can use basic formatting:\n**Bold text**\n*Italic text*\n# Heading\n- Bullet point",
            height=300
        )
        
        if st.button("💾 Save Note", key="save_note", use_container_width=True):
            if note_title and note_content:
                new_note = {
                    "title": note_title,
                    "content": note_content,
                    "category": selected_category,
                    "tags": [tag.strip() for tag in tags_input.split(',')] if tags_input else [],
                    "created": time.time(),
                    "modified": time.time()
                }
                st.session_state.notes.append(new_note)
                
                # Save to local for offline access
                save_to_local(st.session_state.current_user, "notes", new_note)
                
                st.success("✅ Note saved successfully!")
                st.rerun()
            else:
                st.warning("Please enter both title and content")
        
        # Display existing notes
        if st.session_state.notes:
            st.markdown("### 📚 Your Notes")
            for i, note in enumerate(st.session_state.notes):
                with st.expander(f"📝 {note['title']} ({note['category']})"):
                    st.markdown(f"**Created:** {time.ctime(note['created'])}")
                    st.markdown(f"**Tags:** {', '.join(note['tags']) if note['tags'] else 'None'}")
                    st.markdown("**Content:**")
                    st.markdown(note['content'])
                    
                    col_del, col_edit = st.columns(2)
                    with col_del:
                        if st.button("🗑️ Delete", key=f"del_note_{i}"):
                            st.session_state.notes.pop(i)
                            st.rerun()
                    with col_edit:
                        if st.button("✏️ Edit", key=f"edit_note_{i}"):
                            st.session_state.editing_note = i
                            st.rerun()

# ==========================
# ADMIN DASHBOARD FEATURE
# ==========================
def show_admin_feature():
    """Display admin dashboard for system management"""
    st.markdown(
        """
        <div style="text-align: center; margin-bottom: 30px;">
            <h1 style="color: #4e54c8; font-size: 32px; font-weight: 700; margin-bottom: 10px; font-family: 'Georgia', serif;">
                👨‍💼 Admin Dashboard
            </h1>
            <p style="color: #666; font-size: 16px; font-family: 'Georgia', serif;">
                System management and analytics for administrators
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Check if user is admin
    if st.session_state.current_user != "admin":
        st.warning("🔒 Admin access required. Please login as admin.")
        return
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("### 📊 System Stats")
        st.metric("Total Users", "1,234")
        st.metric("Active Sessions", "45")
        st.metric("Storage Used", "2.3 GB")
    
    with col2:
        st.markdown("### 🎯 Feature Usage")
        st.metric("Chatbot Queries", "5,678")
        st.metric("Quiz Generated", "234")
        st.metric("Recordings Made", "89")
    
    with col3:
        st.markdown("### ⚙️ System Health")
        st.metric("API Status", "✅ Online")
        st.metric("Database", "✅ Connected")
        st.metric("Storage", "✅ Available")
    
    st.markdown("---")
    
    # System management
    st.markdown("### 🔧 System Management")
    
    tab1, tab2, tab3 = st.tabs(["Users", "Settings", "Logs"])
    
    with tab1:
        st.markdown("#### 👥 User Management")
        st.dataframe({
            "Username": ["admin", "student1", "student2"],
            "Role": ["Admin", "Student", "Student"],
            "Last Login": ["2024-01-15", "2024-01-14", "2024-01-13"],
            "Status": ["Active", "Active", "Inactive"]
        })
    
    with tab2:
        st.markdown("#### ⚙️ System Settings")
        st.checkbox("Enable Registration", value=True)
        st.checkbox("Maintenance Mode", value=False)
        st.slider("Max File Size (MB)", 1, 100, 10)
    
    with tab3:
        st.markdown("#### 📋 System Logs")
        st.text_area("Recent Logs", value="2024-01-15 10:30:15 - User login: admin\n2024-01-15 10:25:32 - Quiz generated: student1\n2024-01-15 10:20:45 - Recording started: student2", height=200)

# ==========================
# CATEGORIZED SEARCH FEATURE
# ==========================
def show_search_feature():
    """Display categorized search functionality"""
    st.markdown(
        """
        <div class="welcome-container">
            <h1 class="welcome-title">🔍 Categorized Search</h1>
            <p class="welcome-text">
                Search across all your content with advanced filtering and categorization
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### 🔍 Search Filters")
        st.markdown("---")
        
        # Search categories
        search_categories = st.multiselect(
            "Search in:",
            ["Chat History", "Notes", "Quizzes", "Recordings", "Flash Cards"],
            default=["Chat History", "Notes"]
        )
        
        # Date range
        date_range = st.date_input("Date Range", value=[])
        
        # Content type
        content_type = st.selectbox("Content Type", ["All", "Text", "Audio", "Images", "Documents"])
        
        # Tags filter
        tags_filter = st.text_input("Tags", placeholder="e.g., important, exam")
    
    with col2:
        st.markdown("### Search Query")
        
        search_query = st.text_input(
            "Enter your search query",
            placeholder="Search for specific topics, keywords, or phrases...",
            key="search_input"
        )
        
        if st.button(" Search", key="perform_search", use_container_width=True):
            if search_query:
                with st.spinner("Searching..."):
                    results = perform_search(search_query, search_categories, date_range, content_type, tags_filter)
                    st.session_state.search_results = results
                    st.success(f"✅ Found {len(results)} results!")
            else:
                st.warning("Please enter a search query")
        
        # Display search results
        if st.session_state.get('search_results'):
            st.markdown("###  Search Results")
            display_search_results(st.session_state.search_results)

def perform_search(query, categories, date_range, content_type, tags_filter):
    """Perform SEARCH on real saved data from local storage."""
    results = []
    
    if not st.session_state.get('authenticated'):
        return []
        
    username = st.session_state.current_user
    
    # Map UI category labels to folder names
    category_map = {
        "Quizzes": "quizzes",
        "Notes": "notes",
        "Recordings": "transcripts",
        "Flash Cards": "flashcards",
        "Chat History": "chat" # Placeholder even if empty
    }
    
    query_lower = query.lower()
    
    for ui_cat in categories:
        folder_name = category_map.get(ui_cat)
        if not folder_name:
            continue
            
        # Load real items from local storage
        items = load_from_local(username, folder_name)
        
        for item in items:
            # Extract basic info
            title = ""
            content = ""
            tags = item.get('tags', [])
            saved_at = item.get('saved_at', 0)
            date_str = time.strftime('%Y-%m-%d', time.localtime(saved_at)) if saved_at else "Unknown"
            
            if folder_name == "quizzes":
                title = item.get('title') or item.get('subject') or "Untitled Quiz"
                content = item.get('content', '')
            elif folder_name == "notes":
                title = item.get('title', 'Untitled Note')
                content = item.get('content', '')
            elif folder_name == "flashcards":
                title = item.get('subject', 'Untitled Deck')
                # For flashcards, search through questions and answers
                cards = item.get('cards', [])
                content = " ".join([f"{c.get('front')} {c.get('back')}" for c in cards])
            elif folder_name == "transcripts":
                title = item.get('title', 'Untitled Recording')
                content = item.get('content', '')
                
            # Perform search check
            if query_lower in title.lower() or query_lower in content.lower():
                # Simple relevance calculation
                relevance = 1.0
                if query_lower in title.lower():
                    relevance += 0.5 # Title matches are more relevant
                
                results.append({
                    "title": title,
                    "content": content,
                    "category": ui_cat,
                    "date": date_str,
                    "tags": tags if tags else [ui_cat.lower()],
                    "relevance": relevance
                })
    
    return sorted(results, key=lambda x: x['relevance'], reverse=True)

def display_search_results(results):
    """Display search results in a formatted way"""
    for i, result in enumerate(results):
        with st.expander(f"📄 {result['title']} (Relevance: {result['relevance']:.2f})"):
            st.markdown(f"**Category:** {result['category']}")
            st.markdown(f"**Date:** {result['date']}")
            st.markdown(f"**Tags:** {', '.join(result['tags'])}")
            st.markdown(f"**Content:** {result['content'][:200]}...")
            
            if st.button("📖 View Full", key=f"view_result_{i}"):
                st.markdown(f"**Full Content:**\n\n{result['content']}")

# ==========================
# OFFLINE MODE FEATURE
# ==========================
def show_offline_feature():
    """Display offline mode functionality with access to saved content"""
    # 1. Header & Status
    st.markdown(
        """
        <div class="welcome-container">
            <h1 class="welcome-title">Offline Mode</h1>
            <p class="welcome-text">
                Access your previously generated content without internet
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Check connection
    try:
        requests.get("https://www.google.com", timeout=2)
        is_online = True
        status_color = "#28a745"
        status_text = "Connected"
    except:
        is_online = False
        status_color = "#dc3545"
        status_text = "Disconnected"
        
    st.markdown(
        f"""
        <div style="background: white; padding: 15px; border-radius: 12px; border-left: 5px solid {status_color}; box-shadow: 0 2px 10px rgba(0,0,0,0.05); margin-bottom: 25px; display: flex; align-items: center; justify-content: space-between;">
            <div>
                <strong style="color: {status_color}; font-size: 18px;">● {status_text}</strong>
                <div style="color: #666; font-size: 14px; margin-top: 5px;">
                    {'You have full access to all features.' if is_online else 'Restricted to saved content only.'}
                </div>
            </div>
            <div style="text-align: right;">
                <span style="background: #f0f2f6; padding: 5px 10px; border-radius: 15px; font-size: 12px; color: #555;">
                    Data Source: Local Storage
                </span>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    if not is_online:
        st.info("⚠️ You are currently offline. AI features (Chat, Quiz Generation) are unavailable. displaying locally saved copies.")
        
    # 2. Offline Content Viewer
    if not st.session_state.current_user:
        st.warning("Please log in to view saved content.")
        return

    st.markdown("### 📂 Saved Library")
    
    tab1, tab2, tab3, tab4 = st.tabs(["📝 Saved Quizzes", "🎴 Flash Cards", "📔 Notes", "🎙️ Transcripts"])
    
    # --- Saved Quizzes ---
    with tab1:
        quizzes = load_from_local(st.session_state.current_user, "quizzes")
        if not quizzes:
            st.info("No saved quizzes found. Generate quizzes while online to see them here.")
        else:
            for q in quizzes:
                date_str = time.strftime('%Y-%m-%d %H:%M', time.localtime(q.get('saved_at', 0)))
                with st.expander(f"📝 Quiz: {q.get('subject', 'General')} - {date_str}"):
                    st.markdown(f"**Difficulty:** {q.get('difficulty')} | **Questions:** {q.get('num_questions')}")
                    if st.button("👁️ View Quiz", key=f"view_quiz_{q.get('saved_at')}"):
                        st.markdown("---")
                        st.text(q.get('content', 'No content'))
    
    # --- Saved Flash Cards ---
    with tab2:
        flashcards_sets = load_from_local(st.session_state.current_user, "flashcards")
        if not flashcards_sets:
            st.info("No saved flash cards found.")
        else:
            for fc_set in flashcards_sets:
                date_str = time.strftime('%Y-%m-%d %H:%M', time.localtime(fc_set.get('saved_at', 0)))
                with st.expander(f"🎴 Deck: {fc_set.get('subject', 'General')} - {date_str}"):
                    cards = fc_set.get('cards', [])
                    st.markdown(f"**{len(cards)} Cards**")
                    if st.button("👁️ View Deck", key=f"view_fc_{fc_set.get('saved_at')}"):
                        for card in cards:
                            st.markdown(f"**Q:** {card['front']}")
                            st.markdown(f"**A:** {card['back']}")
                            st.markdown("---")

    # --- Saved Notes ---
    with tab3:
        notes = load_from_local(st.session_state.current_user, "notes")
        if not notes:
            st.info("No saved notes found.")
        else:
            for note in notes:
                date_str = time.strftime('%Y-%m-%d %H:%M', time.localtime(note.get('saved_at', 0)))
                with st.expander(f"📔 {note.get('title', 'Untitled')} - {date_str}"):
                    st.markdown(f"**Category:** {note.get('category')}")
                    st.markdown("---")
                    st.markdown(note.get('content', ''))

    # --- Saved Transcripts ---
    with tab4:
        transcripts = load_from_local(st.session_state.current_user, "transcripts")
        if not transcripts:
            st.info("No saved transcripts found.")
        else:
            for trans in transcripts:
                date_str = time.strftime('%Y-%m-%d %H:%M', time.localtime(trans.get('saved_at', 0)))
                with st.expander(f"🎙️ Recording: {trans.get('title', 'Untitled')} - {date_str}"):
                    st.text_area("Transcript", trans.get('content', ''), height=300, key=f"read_trans_{trans.get('saved_at')}")

def show_profile_feature():
    """Display user profile feature with modern, minimal UI"""
    st.markdown(
        """
        <div class="welcome-container">
            <h1 class="welcome-title">👤 User Profile</h1>
            <p class="welcome-text">Manage your personal identity, view insights, and control your account.</p>
        </div>
        """,
        unsafe_allow_html=True
    )

    if not st.session_state.user_id:
        st.warning("Please log in to view your profile.")
        return

    # Fetch latest user data (for display name, etc.)
    user_data = get_full_user_data(st.session_state.user_id)
    profile = user_data.get("profile", {})
    stats = user_data.get("stats", {})
    
    # Sync display name to session state for navigation
    if profile.get("display_name"):
        st.session_state.user_display_name = profile["display_name"]
    
    # Custom CSS for Profile Cards
    st.markdown("""
        <style>
        
        .profile-label {
            font-size: 14px;
            color: #6b7280;
            margin-bottom: 5px;
            font-weight: 600;
        }
        .profile-value {
            font-size: 16px;
            color: #111827;
            font-weight: 500;
            margin-bottom: 15px;
        }
        .insight-card {
            background: #f9fafb;
            padding: 20px;
            border-radius: 12px;
            border: 1px solid #f3f4f6;
            text-align: center;
        }
        .danger-zone {
            border: 1px solid #fee2e2;
            background: #fffafb;
            padding: 20px;
            border-radius: 12px;
        }
        </style>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3 = st.tabs(["👤 Identity", "📈 Insights", "🔐 Settings"])

    with tab1:
        st.markdown('<div class="profile-section">', unsafe_allow_html=True)
        col1, col2 = st.columns([1, 3])
        
        with col1:
            # Simple circular avatar placeholder
            initials = profile.get("username", "U")[:2].upper()
            st.markdown(f"""
                <div style="
                    width: 100px; height: 100px; 
                    background: linear-gradient(135deg, #4e54c8, #8f94fb); 
                    border-radius: 50%; 
                    display: flex; align-items: center; justify-content: center; 
                    color: white; font-size: 32px; font-weight: 700;
                    margin: 0 auto 20px auto;
                ">
                    {initials}
                </div>
            """, unsafe_allow_html=True)
            
        with col2:
            current_display_name = profile.get("display_name") or profile.get("username")
            st.markdown(f'<div class="profile-label">Username</div><div class="profile-value">{profile.get("username")}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="profile-label">Email Address</div><div class="profile-value">{profile.get("email") or "Not provided"}</div>', unsafe_allow_html=True)
            
            new_display_name = st.text_input("Display Name", value=current_display_name, help="How you want to be called in the app")
            
            if st.button("Update Profile", type="primary"):
                success, msg = update_user_profile(st.session_state.user_id, display_name=new_display_name)
                if success:
                    st.success("Profile updated! Refreshing...")
                    st.session_state.user_display_name = new_display_name
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error(msg)
        st.markdown('</div>', unsafe_allow_html=True)

    with tab2:
        st.markdown('<div class="profile-section">', unsafe_allow_html=True)
        st.subheader("Light Usage Insights")
        st.markdown("These insights reflect your journey with Lecturebuddies. Keep up the great work!")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown('<div class="insight-card">', unsafe_allow_html=True)
            st.metric("Study Sessions", stats.get("study_sessions", 0))
            st.markdown('</div>', unsafe_allow_html=True)
            
        with col2:
            st.markdown('<div class="insight-card">', unsafe_allow_html=True)
            st.metric("Quizzes Created", stats.get("quizzes_created", 0))
            st.markdown('</div>', unsafe_allow_html=True)
            
        with col3:
            st.markdown('<div class="insight-card">', unsafe_allow_html=True)
            st.metric("Recordings", stats.get("recordings_count", 0))
            st.markdown('</div>', unsafe_allow_html=True)
            
        st.markdown("---")
        
        # Calculate most used feature
        feature_counts = {
            "Chatbot": stats.get("study_sessions", 0),
            "Quizzes": stats.get("quizzes_created", 0),
            "Recordings": stats.get("recordings_count", 0),
            "Flashcards": stats.get("flashcards_created", 0),
            "Notes": stats.get("notes_count", 0)
        }
        most_used = max(feature_counts, key=feature_counts.get)
        
        st.info(f"✨ **Motivational Tip:** You mostly use **{most_used}**. Try exploring other features to diversify your study routine!")
        st.markdown('</div>', unsafe_allow_html=True)

    with tab3:
        st.markdown('<div class="profile-section">', unsafe_allow_html=True)
        st.subheader("Account Management")
        
        with st.expander("🔐 Change Password"):
            new_pwd = st.text_input("New Password", type="password")
            confirm_pwd = st.text_input("Confirm New Password", type="password")
            if st.button("Change Password"):
                if new_pwd == confirm_pwd and len(new_pwd) >= 6:
                    success, msg = update_user_profile(st.session_state.user_id, new_password=new_pwd)
                    if success: st.success("Password changed successfully!")
                    else: st.error(msg)
                else:
                    st.error("Passwords must match and be at least 6 characters.")
        
        st.markdown("---")
        
        col_export, col_logout = st.columns(2)
        with col_export:
            st.markdown("**Data Portability**")
            st.markdown("Download a copy of all your data.")
            if st.button("Export Data (JSON)"):
                full_data = get_full_user_data(st.session_state.user_id)
                st.download_button(
                    label="Download JSON",
                    data=json.dumps(full_data, indent=2),
                    file_name=f"lecturebuddies_data_{profile.get('username')}.json",
                    mime="application/json"
                )
                
        with col_logout:
            st.markdown("**Session Control**")
            st.markdown("Sign out of your account.")
            if st.button("🚪 Logout Now", key="profile_logout_btn"):
                st.session_state.authenticated = False
                st.session_state.current_user = None
                st.session_state.user_id = None
                st.session_state.selected_feature = None
                st.rerun()

        st.markdown("---")
        st.markdown('<div class="danger-zone">', unsafe_allow_html=True)
        st.markdown('<h4 style="color: #c53030; margin-top: 0;">⚠️ Danger Zone</h4>', unsafe_allow_html=True)
        st.markdown("Deleting your account is permanent and cannot be undone.")
        
        if "confirm_delete" not in st.session_state:
            st.session_state.confirm_delete = False
            
        if not st.session_state.confirm_delete:
            if st.button("Delete My Account", type="secondary"):
                st.session_state.confirm_delete = True
                st.rerun()
        else:
            st.warning("Are you absolutely sure? This will delete all your quizzes, notes, and recordings.")
            col_d1, col_d2 = st.columns(2)
            with col_d1:
                if st.button("🔥 Yes, Delete Permanently", type="primary"):
                    success, msg = delete_user_account(st.session_state.user_id)
                    if success:
                        st.session_state.authenticated = False
                        st.session_state.current_user = None
                        st.session_state.user_id = None
                        st.session_state.selected_feature = None
                        st.session_state.confirm_delete = False
                        st.success("Account deleted. Goodbye!")
                        time.sleep(2)
                        st.rerun()
                    else:
                        st.error(msg)
            with col_d2:
                if st.button("Cancel"):
                    st.session_state.confirm_delete = False
                    st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

# ==========================
# HELPER FUNCTIONS
# ==========================
def inject_file_content(user_message: str) -> str:
    """Replace file references in user message with extracted text"""
    for fname, content in st.session_state.document_contents.items():
        if fname.lower() in user_message.lower():
            extracted = content if content.strip() else "[No text extracted from this file]"
            user_message = user_message.replace(
                fname,
                f"(Extracted content: {extracted[:1000]}...)"
            )
    return user_message

def get_groq_response(user_input, model="llama-3.1-8b-instant", temperature=0.7):
    """Send query + context to Groq API and return assistant response"""
    if not api_key or api_key.strip() == "":
        return "Missing API key. Please set GROQ_API_KEY in your .env file."

    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    # Build document context
    doc_context = ""
    if st.session_state.document_contents:
        doc_context = "\n\n**Available Documents:**\n"
        for fname, content in st.session_state.document_contents.items():
            snippet = content[:1000] if content else "[No extractable text]"
            doc_context += f"\n--- {fname} ---\n{snippet}...\n"

    # Build conversation
    messages = [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages]
    system_msg = (
        f"You are Lecturebuddies, an AI chatbot designed for education. "
        f"Answer clearly, summarize effectively, and explain concepts step by step."
        f"{' Available Documents: ' + doc_context if doc_context else ''}\n\n"
        "Guidelines:\n"
        "Education-focused\n"
        "Summarization expert\n"
        "Clarity first (simple language, then details)\n"
        "Confidence + accuracy\n"
        "Break down topics step-by-step, use examples, and stay professional yet supportive."
    )
    messages.insert(0, {"role": "system", "content": system_msg})
    messages.append({"role": "user", "content": user_input})

    payload = {"model": model, "messages": messages, "temperature": float(temperature), "max_tokens": 1000}

    try:
        resp = requests.post(url, headers=headers, json=payload, timeout=30)
        if resp.status_code == 200:
            data = resp.json()
            return data.get("choices", [{}])[0].get("message", {}).get("content", "No response received.")
        elif resp.status_code == 401:
            return "Invalid API key. Please check GROQ_API_KEY in your .env file."
        elif resp.status_code == 429:
            return "Too many requests. Please slow down and retry shortly."
        else:
            return f"API Error {resp.status_code}: {resp.text}"
    except requests.exceptions.Timeout:
        return "Request timed out. Please retry."
    except requests.exceptions.RequestException as e:
        return f"Network error: {e}"
    except Exception as e:
        return f"Unexpected error: {e}"

def process_document(uploaded_file, doc_processor):
    """Extract text from uploaded documents"""
    try:
        temp_path = os.path.join("temp", uploaded_file.name)
        os.makedirs("temp", exist_ok=True)
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.read())

        content = doc_processor.process_document(temp_path)
        return content if content.strip() else "[No text extracted]"
    except Exception as e:
        return f"[File processing error: {e}]"

def get_groq_quiz_response(content, num_questions=5, difficulty="Medium", model="llama-3.1-8b-instant", temperature=0.7):
    """Send content to Groq API and get quiz questions back"""
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    difficulty_map = {"Easy": "simple and straightforward", "Medium": "balanced and informative", "Hard": "challenging and detailed"}
    diff_desc = difficulty_map.get(difficulty, "balanced")

    system_msg = (
        "You are Lecturebuddies Quiz Generator. "
        "Your task is to generate high-quality multiple-choice quizzes (MCQs) from educational material. "
        f"Generate exactly {num_questions} MCQs. Each question should follow this strict structure:\n"
        "1. **Question Text**\n\n"
        "A) Option 1\n\n"
        "B) Option 2\n\n"
        "C) Option 3\n\n"
        "D) Option 4\n\n"
        "Correct: [Letter]\n\n"
        "CRITICAL: Use DOUBLE NEW LINES between each option and the question to ensure they appear on separate lines in Markdown rendering. "
        f"Make questions {diff_desc} in difficulty. "
        "Format the output neatly with numbered questions and bold question text. "
        "End with a summary of correct answers."
    )

    messages = [
        {"role": "system", "content": system_msg},
        {"role": "user", "content": f"Generate a quiz from this content:\n\n{content}"}
    ]

    payload = {"model": model, "messages": messages, "temperature": float(temperature), "max_tokens": 1200}

    try:
        with st.spinner(f"Generating {num_questions} {difficulty} quiz questions..."):
            resp = requests.post(url, headers=headers, json=payload, timeout=30)
        if resp.status_code == 200:
            data = resp.json()
            raw_content = data.get("choices", [{}])[0].get("message", {}).get("content", "No response received.")
            
            # Post-processing to ENSURE newlines before options if the model misses them
            import re
            for opt in ["A\)", "B\)", "C\)", "D\)", "Correct:"]:
                # Look for the option and ensure it has at least two newlines before it
                raw_content = re.sub(rf"(?<!\n\n)\s*({opt})", r"\n\n\1", raw_content)
            
            return raw_content
        elif resp.status_code == 401:
            return "Invalid API key. Please check GROQ_API_KEY in your .env file."
        elif resp.status_code == 429:
            return "Too many requests. Please retry shortly."
        else:
            return f"API Error {resp.status_code}: {resp.text}"
    except requests.exceptions.Timeout:
        return "Request timed out. Please retry."
    except Exception as e:
        return f"Error: {str(e)}"

def extract_content_from_file(uploaded_file):
    """Extract text content from uploaded file (PDF, DOCX, TXT)"""
    file_type = uploaded_file.type if hasattr(uploaded_file, 'type') else uploaded_file.name.split('.')[-1].lower()

    try:
        if file_type == "application/pdf" or file_type.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            content = ""
            for page in pdf_reader.pages:
                text = page.extract_text() or ""
                content += text + "\n"
            return content.strip()
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_type.endswith('.docx'):
            doc = Document(io.BytesIO(uploaded_file.read()))
            content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            return content.strip()
        else:  # TXT or other text-based
            return uploaded_file.read().decode("utf-8", errors="ignore").strip()
    except Exception as e:
        return f"Error extracting content: {str(e)}"

def save_to_local(username, category, data, custom_filename=None):
    """Save content to local JSON for offline access."""
    if not username:
        return
    
    base_dir = "user_data"
    user_dir = os.path.join(base_dir, username)
    category_dir = os.path.join(user_dir, category)
    os.makedirs(category_dir, exist_ok=True)
    
    if custom_filename:
        filename = custom_filename
        if not filename.endswith(".json"):
            filename += ".json"
    else:
        # Create a unique filename based on timestamp
        title_part = ""
        if isinstance(data, dict):
            if "title" in data:
                title_part = f"_{data['title'][:20].replace(' ', '_')}"
            elif "subject" in data:
                title_part = f"_{data['subject'][:20].replace(' ', '_')}"
        
        # Clean filename
        title_part = "".join([c for c in title_part if c.isalnum() or c in ('_', '-')])
        
        filename = f"{int(time.time())}{title_part}_{category}.json"
        
    file_path = os.path.join(category_dir, filename)
    
    # Add timestamp
    if isinstance(data, dict) and 'saved_at' not in data:
        data['saved_at'] = int(time.time())
    
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Error saving to local: {e}")

def load_from_local(username, category):
    """Load all content from local JSON for a category."""
    if not username:
        return []
        
    base_dir = "user_data"
    category_dir = os.path.join(base_dir, username, category)
    
    if not os.path.exists(category_dir):
        return []
        
    items = []
    try:
        for filename in os.listdir(category_dir):
            if filename.endswith(".json"):
                file_path = os.path.join(category_dir, filename)
                with open(file_path, "r", encoding="utf-8") as f:
                    try:
                        data = json.load(f)
                        if isinstance(data, dict) and 'saved_at' not in data:
                            try:
                                ts = int(filename.split('_')[0])
                                data['saved_at'] = ts
                            except:
                                data['saved_at'] = 0
                        items.append(data)
                    except json.JSONDecodeError:
                        continue
    except Exception as e:
        print(f"Error loading from local: {e}")
        
    return sorted(items, key=lambda x: x.get('saved_at', 0), reverse=True)

# Recording helper functions
def record_worker(audio_q: queue.Queue, stop_event: threading.Event):
    """Recording thread worker"""
    while not stop_event.is_set():
        time.sleep(0.1)

def sd_callback(indata, frames, time_info, status):
    """Callback for sounddevice.InputStream"""
    if status:
        pass
    mono = np.mean(indata, axis=1).astype(np.float32)
    # guard: audio_queue may be None if recording not properly started
    if st.session_state.get("audio_queue") is not None:
        try:
            st.session_state.audio_queue.put(mono)
        except Exception:
            pass

def transcription_consumer(model: WhisperModel, chunk_seconds=5):
    """Transcription consumer thread"""
    sample_per_chunk = chunk_seconds * 16000
    buffer = np.zeros((0,), dtype=np.float32)
    chunk_index = 0

    while st.session_state.recording or (st.session_state.audio_queue is not None and not st.session_state.audio_queue.empty()):
        try:
            while st.session_state.audio_queue is not None and not st.session_state.audio_queue.empty():
                data = st.session_state.audio_queue.get_nowait()
                buffer = np.concatenate((buffer, data))

            if buffer.shape[0] >= sample_per_chunk:
                to_process = buffer[:sample_per_chunk]
                buffer = buffer[sample_per_chunk:]

                with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
                    tmp_path = tmp.name
                    sf.write(tmp_path, to_process, 16000, subtype="PCM_16")

                segments, info = model.transcribe(tmp_path, beam_size=5)
                partial_text = ""
                for seg in segments:
                    partial_text += seg.text

                prev = st.session_state.transcript
                st.session_state.transcript = (prev + " " + partial_text).strip()
                st.session_state.partial_transcript = partial_text

                saved_chunk = os.path.join("temp_recordings", f"chunk_{int(time.time())}_{chunk_index}.wav")
                os.replace(tmp_path, saved_chunk)
                st.session_state.chunks_saved.append(saved_chunk)
                chunk_index += 1
            else:
                time.sleep(0.2)
        except Exception:
            time.sleep(0.1)
    return

# Realtime recording helper functions
def realtime_record_worker(audio_q: queue.Queue, stop_event: threading.Event):
    """Recording thread worker for realtime"""
    while not stop_event.is_set():
        time.sleep(0.1)

def realtime_sd_callback(indata, frames, time_info, status):
    """Callback for sounddevice.InputStream in realtime mode"""
    if status:
        pass
    mono = np.mean(indata, axis=1).astype(np.float32)
    if st.session_state.get("realtime_queue") is not None:
        try:
            st.session_state.realtime_queue.put(mono)
        except Exception:
            pass

def realtime_transcription_consumer(model: WhisperModel, chunk_seconds=5):
    """Transcription consumer thread for realtime"""
    sample_per_chunk = chunk_seconds * 16000
    buffer = np.zeros((0,), dtype=np.float32)
    chunk_index = 0

    while st.session_state.realtime_recording or (st.session_state.realtime_queue is not None and not st.session_state.realtime_queue.empty()):
        try:
            # Get audio data from queue
            while st.session_state.realtime_queue is not None and not st.session_state.realtime_queue.empty():
                data = st.session_state.realtime_queue.get_nowait()
                buffer = np.concatenate((buffer, data))

            # Process when we have enough data
            if buffer.shape[0] >= sample_per_chunk:
                to_process = buffer[:sample_per_chunk]
                buffer = buffer[sample_per_chunk:]

                # Save audio chunk to temporary file
                with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
                    tmp_path = tmp.name
                    sf.write(tmp_path, to_process, 16000, subtype="PCM_16")

                # Transcribe the chunk
                try:
                    segments, info = model.transcribe(tmp_path, beam_size=5)
                    partial_text = ""
                    for seg in segments:
                        partial_text += seg.text

                    # Update session state
                    if partial_text.strip():
                        prev = st.session_state.realtime_transcript
                        st.session_state.realtime_transcript = (prev + " " + partial_text).strip()
                        st.session_state.realtime_partial = partial_text
                        
                        # Save chunk for debugging
                        saved_chunk = os.path.join("temp_recordings", f"realtime_chunk_{int(time.time())}_{chunk_index}.wav")
                        os.replace(tmp_path, saved_chunk)
                        st.session_state.realtime_chunks.append(saved_chunk)
                        chunk_index += 1
                        
                        print(f"Transcribed chunk {chunk_index}: {partial_text}")  # Debug print
                    else:
                        print(f"No text in chunk {chunk_index}")  # Debug print
                        os.remove(tmp_path)
                        
                except Exception as e:
                    print(f"Transcription error: {e}")  # Debug print
                    try:
                        os.remove(tmp_path)
                    except:
                        pass
            else:
                time.sleep(0.2)
        except Exception as e:
            print(f"Consumer error: {e}")  # Debug print
            time.sleep(0.1)
    return

# ==========================
# MAIN APPLICATION LOGIC
# ==========================
def main():
    """Main application entry point"""
    # Simple Forgot Password routing
    if st.session_state.get("show_simple_reset", False):
        show_simple_password_reset()
        return

    if not st.session_state.authenticated:
        show_login_page()
    else:
        show_dashboard()

if __name__ == "__main__":
    main()
